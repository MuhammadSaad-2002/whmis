#!/usr/bin/env python3
"""Condensed WHMIS commercial deliverables.

Produces two tight, client-facing .docx files that reuse the slate/emerald
corporate styling but are engineered for a fixed page budget:

  1. WHMIS-PreDeployment-Requirements-and-Payment-Schedule-2page.docx  (<= 2 pages)
  2. WHMIS-MOU.docx                                                    (1 page)

    python3 scripts/build-predeployment-condensed.py
"""
import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTDIR = os.path.join(ROOT, "docs", "commercial")
OUT_2P = os.path.join(OUTDIR, "WHMIS-PreDeployment-Requirements-and-Payment-Schedule-2page.docx")
OUT_MOU = os.path.join(OUTDIR, "WHMIS-MOU.docx")
OUT_PW = os.path.join(OUTDIR, "WHMIS-PhaseWise-Development-and-Deployment-Plan.docx")
OUT_INV = os.path.join(OUTDIR, "WHMIS-Invoice.docx")

VENDOR = "Virtual Wisdom Technologies"
CLIENT = "Master Pharmaceuticals Distributor"
PROJECT = "WHMIS"
VERSION = "1.0"
DOC_DATE = "August 2026"
CLASSIFICATION = "Confidential – Commercial & Project Information"

SLATE = RGBColor(0x0F, 0x17, 0x2A)
EMERALD = RGBColor(0x05, 0x96, 0x69)
EMERALD_DK = RGBColor(0x04, 0x78, 0x57)
GREY = RGBColor(0x6B, 0x72, 0x80)
LIGHT = RGBColor(0x9C, 0xA3, 0xAF)
INK = RGBColor(0x1F, 0x29, 0x37)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SLATE_HEX = "0F172A"
EMERALD_HEX = "059669"
CALLOUT_FILL = "ECFDF5"
SNAP_FILL = "F0FDF4"
ZEBRA_FILL = "F9FAFB"
TOTAL_FILL = "E5E7EB"
RULE = "D1D5DB"

# Blue brand theme (#004aad) — alternative to the emerald accent.
BLUE = RGBColor(0x00, 0x4A, 0xAD)
BLUE_DK = RGBColor(0x00, 0x3A, 0x87)
BLUE_HEX = "004AAD"
BLUE_FILL = "E7EFFA"   # light blue tint
BLUE_HAIR = "9DBBE8"   # light blue hairline
BAND_SUB = RGBColor(0xCF, 0xDE, 0xF4)   # light text on the blue band
HAIR = "E5E7EB"        # very light inside-row hairline

# Active accent (swappable per document via set_theme). Defaults to emerald.
A_MAIN = EMERALD
A_DK = EMERALD_DK
A_HEX = EMERALD_HEX
A_FILL = CALLOUT_FILL
A_HAIR = "A7F3D0"


def set_theme(main, dk, hex_, fill, hair):
    global A_MAIN, A_DK, A_HEX, A_FILL, A_HAIR
    A_MAIN, A_DK, A_HEX, A_FILL, A_HAIR = main, dk, hex_, fill, hair


THEME_EMERALD = (EMERALD, EMERALD_DK, EMERALD_HEX, CALLOUT_FILL, "A7F3D0")
THEME_BLUE = (BLUE, BLUE_DK, BLUE_HEX, BLUE_FILL, BLUE_HAIR)


# ----------------------------- low-level ---------------------------------- #
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def add_field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    instr = _el("w:instrText", **{"xml:space": "preserve"})
    instr.text = instruction
    run._r.append(instr)
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "separate"}))
    if placeholder:
        t = OxmlElement("w:t")
        t.text = placeholder
        run._r.append(t)
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))


def _spc(run, pts):
    run._r.get_or_add_rPr().set("spc", str(int(pts * 100)))


def _shade(cell, fill):
    cell._tc.get_or_add_tcPr().append(_el("w:shd", **{"w:val": "clear", "w:fill": fill}))


def _cell_margins(cell, top=30, bottom=30, left=70, right=70):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = _el("w:tcMar")
    for edge, w in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        mar.append(_el(f"w:{edge}", **{"w:w": str(w), "w:type": "dxa"}))
    tcPr.append(mar)


def _borders(el_owner, tag, color=RULE, sz="4", edges=("top", "left", "bottom", "right", "insideH", "insideV")):
    b = _el(tag)
    for edge in edges:
        b.append(_el(f"w:{edge}", **{"w:val": "single", "w:sz": sz, "w:space": "0", "w:color": color}))
    el_owner.append(b)
    return b


def blank(doc, pts=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(pts)
    return p


def set_margins(doc, top, bottom, left, right, section=None):
    s = doc.sections[0] if section is None else section
    s.top_margin = Inches(top)
    s.bottom_margin = Inches(bottom)
    s.left_margin = Inches(left)
    s.right_margin = Inches(right)
    s.header_distance = Inches(0.3)
    s.footer_distance = Inches(0.3)
    return s


def base_styles(doc, body_size, line=1.02):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(body_size)
    st.paragraph_format.space_after = Pt(2)
    st.paragraph_format.line_spacing = line


def _set_cell_text(cell, text, bold=False, color=None, size=8, align=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color


def _run_shade(run, fill):
    run._r.get_or_add_rPr().append(_el("w:shd", **{"w:val": "clear", "w:fill": fill, "w:color": "auto"}))


def _cell_border(cell, edge, sz, color):
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.find(qn("w:tcBorders"))
    if b is None:
        b = _el("w:tcBorders")
        tcPr.append(b)
    b.append(_el(f"w:{edge}", **{"w:val": "single", "w:sz": str(sz), "w:space": "0", "w:color": color}))


def chip_run(p, num, fill=None, size=9):
    fill = A_HEX if fill is None else fill
    r = p.add_run(f"  {num}  ")
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = WHITE
    _run_shade(r, fill)
    gap = p.add_run("  ")
    gap.font.size = Pt(size)
    return r


def title_band(doc, title, subtitle, after=5):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    cell = t.rows[0].cells[0]
    cell.width = Inches(CONTENT_W)
    _shade(cell, A_HEX)
    _cell_margins(cell, top=150, bottom=150, left=170, right=170)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run("WHMIS  ·  PHARMACEUTICAL DISTRIBUTION ERP")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = BAND_SUB
    _spc(r, 1.2)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.0
    r = p2.add_run(title)
    r.bold = True
    r.font.size = Pt(19)
    r.font.color.rgb = WHITE
    _spc(r, 0.3)
    p3 = cell.add_paragraph()
    p3.paragraph_format.space_after = Pt(0)
    p3.paragraph_format.line_spacing = 1.0
    r = p3.add_run(subtitle)
    r.font.size = Pt(8.5)
    r.font.color.rgb = BAND_SUB
    blank(doc, after)
    return t


def table(doc, headers, rows, widths, right_cols=(), total_row=None, size=8, after=3,
          header_fill=None):
    header_fill = A_HEX if header_fill is None else header_fill
    right_cols = set(right_cols)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    # Hairline frame: light top/bottom rules + faint inside-horizontal only.
    tb = _el("w:tblBorders")
    tb.append(_el("w:top", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": RULE}))
    tb.append(_el("w:bottom", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": RULE}))
    tb.append(_el("w:insideH", **{"w:val": "single", "w:sz": "2", "w:space": "0", "w:color": HAIR}))
    t._tbl.tblPr.append(tb)

    def _w(row):
        for cell, w in zip(row.cells, widths):
            cell.width = Inches(w)

    hdr = t.rows[0]
    _w(hdr)
    for i, (cell, text) in enumerate(zip(hdr.cells, headers)):
        _shade(cell, header_fill)
        _cell_margins(cell, top=55, bottom=55, left=95, right=95)
        _set_cell_text(cell, text, bold=True, color=WHITE, size=size,
                       align=WD_ALIGN_PARAGRAPH.RIGHT if i in right_cols else WD_ALIGN_PARAGRAPH.LEFT)
    for ri, row in enumerate(rows):
        r = t.add_row()
        _w(r)
        for i, (cell, text) in enumerate(zip(r.cells, row)):
            if ri % 2 == 1:
                _shade(cell, ZEBRA_FILL)
            _cell_margins(cell, top=48, bottom=48, left=95, right=95)
            _set_cell_text(cell, str(text), size=size, color=INK,
                           align=WD_ALIGN_PARAGRAPH.RIGHT if i in right_cols else WD_ALIGN_PARAGRAPH.LEFT)
    if total_row is not None:
        r = t.add_row()
        _w(r)
        for i, (cell, text) in enumerate(zip(r.cells, total_row)):
            _shade(cell, A_FILL)
            _cell_margins(cell, top=55, bottom=55, left=95, right=95)
            _cell_border(cell, "top", 12, A_HEX)
            _set_cell_text(cell, str(text), bold=True, size=size, color=A_DK,
                           align=WD_ALIGN_PARAGRAPH.RIGHT if i in right_cols else WD_ALIGN_PARAGRAPH.LEFT)
    blank(doc, after)
    return t


def sec(doc, num, title, parts, after=3, size=None):
    """Run-in section: [chip N] bold Title — then inline body parts (text, {opts})."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if num != "":
        p.paragraph_format.space_before = Pt(4)
        chip_run(p, num, size=(size or 9))
        if title:
            rt = p.add_run(title + " — ")
            rt.bold = True
            rt.font.color.rgb = SLATE
            if size:
                rt.font.size = Pt(size)
    else:
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.left_indent = Inches(0.02)
    for text, opts in parts:
        rr = p.add_run(text)
        if opts.get("bold"):
            rr.bold = True
        if opts.get("italic"):
            rr.italic = True
        if "color" in opts:
            rr.font.color.rgb = opts["color"]
        if size:
            rr.font.size = Pt(size)
    return p


def sectitle(doc, num, title, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(6)
    chip_run(p, num, size=size)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = SLATE
    pbdr = _el("w:pBdr")
    pbdr.append(_el("w:bottom", **{"w:val": "single", "w:sz": "4", "w:space": "3", "w:color": A_HEX}))
    p._p.get_or_add_pPr().append(pbdr)
    return p


def callout(doc, title, body, fill=None, accent=None, title_color=None,
            size=8, after=3):
    fill = A_FILL if fill is None else fill
    accent = A_HEX if accent is None else accent
    title_color = A_DK if title_color is None else title_color
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    cell = t.rows[0].cells[0]
    cell.width = Inches(CONTENT_W)
    _shade(cell, fill)
    tcPr = cell._tc.get_or_add_tcPr()
    b = _el("w:tcBorders")
    b.append(_el("w:left", **{"w:val": "single", "w:sz": "18", "w:space": "0", "w:color": accent}))
    for edge in ("top", "bottom", "right"):
        b.append(_el(f"w:{edge}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": A_HAIR}))
    tcPr.append(b)
    _cell_margins(cell, top=50, bottom=50, left=90, right=90)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.02
    if title:
        r = p.add_run(title + "  ")
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = title_color
    r = p.add_run(body)
    r.font.size = Pt(size)
    r.font.color.rgb = INK
    blank(doc, after)
    return t


def two_col(doc, left_title, left_items, right_title, right_items, size=8, after=3):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    _borders(t._tbl.tblPr, "w:tblBorders")
    half = CONTENT_W / 2
    for cell, title, items in ((t.rows[0].cells[0], left_title, left_items),
                               (t.rows[0].cells[1], right_title, right_items)):
        cell.width = Inches(half)
        _cell_margins(cell, top=50, bottom=50, left=90, right=90)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(size + 0.5)
        r.font.color.rgb = A_DK
        for it in items:
            pp = cell.add_paragraph()
            pp.paragraph_format.space_after = Pt(1)
            pp.paragraph_format.line_spacing = 1.0
            rr = pp.add_run("•  " + it)
            rr.font.size = Pt(size)
            rr.font.color.rgb = INK
    blank(doc, after)
    return t


def masthead(doc, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(CLASSIFICATION.upper())
    r.bold = True
    r.font.size = Pt(7.5)
    r.font.color.rgb = A_MAIN
    _spc(r, 0.8)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("WHMIS System — ")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = SLATE
    r = p.add_run(subtitle)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = A_DK

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"Project / System: {PROJECT}     ·     Version {VERSION}     ·     {DOC_DATE}")
    r.font.size = Pt(8.5)
    r.font.color.rgb = GREY
    pbdr = _el("w:pBdr")
    pbdr.append(_el("w:bottom", **{"w:val": "single", "w:sz": "10", "w:space": "3", "w:color": A_HEX}))
    p._p.get_or_add_pPr().append(pbdr)


def parties_row(doc, size=8, after=4):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    _borders(t._tbl.tblPr, "w:tblBorders")
    half = CONTENT_W / 2
    data = [
        ("Development / Technology Partner", VENDOR,
         [("Contact: ", "Muhammad Saad Mubeen, Project Manager"),
          ("Email: ", "pm@vwisdomtechnologies.com"),
          ("Phone: ", "+92 318 5161571")]),
        ("Client", CLIENT,
         [("Representative: ", "Mr. Jamil Ahmed"),
          ("Designation / Email: ", "[Insert]"),
          ("Phone / Address: ", "[Insert]")]),
    ]
    for cell, (role, name, lines) in zip(t.rows[0].cells, data):
        cell.width = Inches(half)
        _cell_margins(cell, top=85, bottom=70, left=120, right=110)
        # accent top bar + hairline sides/bottom = card look
        _cell_border(cell, "top", 22, A_HEX)
        for edge in ("bottom", "left", "right"):
            _cell_border(cell, edge, 4, RULE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(role.upper())
        r.bold = True
        r.font.size = Pt(7)
        r.font.color.rgb = A_MAIN
        _spc(r, 0.8)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.0
        r = p2.add_run(name)
        r.bold = True
        r.font.size = Pt(size + 3)
        r.font.color.rgb = SLATE
        for label, value in lines:
            pl = cell.add_paragraph()
            pl.paragraph_format.space_after = Pt(1)
            pl.paragraph_format.line_spacing = 1.0
            rl = pl.add_run(label)
            rl.font.size = Pt(size - 0.5)
            rl.font.color.rgb = GREY
            rv = pl.add_run(value)
            rv.font.size = Pt(size - 0.5)
            rv.font.color.rgb = INK
    blank(doc, after)


def sign_row(doc, size=8, after=0):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    _borders(t._tbl.tblPr, "w:tblBorders")
    half = CONTENT_W / 2
    data = [(VENDOR, "Muhammad Azhar", "[Insert]"),
            (CLIENT, "[Insert]", "[Insert]")]
    for cell, (org, name, desig) in zip(t.rows[0].cells, data):
        cell.width = Inches(half)
        _cell_margins(cell, top=80, bottom=95, left=120, right=110)
        _cell_border(cell, "top", 22, A_HEX)
        for edge in ("bottom", "left", "right"):
            _cell_border(cell, edge, 4, RULE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run("AUTHORIZED SIGNATORY")
        r.bold = True
        r.font.size = Pt(7)
        r.font.color.rgb = A_MAIN
        _spc(r, 0.8)
        p0 = cell.add_paragraph()
        p0.paragraph_format.space_after = Pt(7)
        p0.paragraph_format.line_spacing = 1.0
        r = p0.add_run("For " + org)
        r.bold = True
        r.font.size = Pt(size + 1)
        r.font.color.rgb = SLATE
        for label, value in (("Name: ", name), ("Designation: ", desig),
                             ("Signature: ", "______________________"),
                             ("Date: ", "______________________")):
            pl = cell.add_paragraph()
            pl.paragraph_format.space_after = Pt(4)
            pl.paragraph_format.line_spacing = 1.0
            rl = pl.add_run(label)
            rl.bold = True
            rl.font.size = Pt(size)
            rl.font.color.rgb = GREY
            rv = pl.add_run(value)
            rv.font.size = Pt(size)
            rv.font.color.rgb = INK
    blank(doc, after)


def footer(doc, text, section=None, page=True):
    sec_obj = doc.sections[0] if section is None else section
    if section is not None:
        sec_obj.footer.is_linked_to_previous = False
    fp = sec_obj.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pbdr = _el("w:pBdr")
    pbdr.append(_el("w:top", **{"w:val": "single", "w:sz": "4", "w:space": "4", "w:color": RULE}))
    fp._p.get_or_add_pPr().append(pbdr)
    r = fp.add_run(text + ("     ·     Page " if page else ""))
    r.font.size = Pt(7.5)
    r.font.color.rgb = GREY
    if page:
        add_field(fp, "PAGE")
    for run in fp.runs:
        run.font.size = Pt(7.5)
        run.font.color.rgb = GREY


def _center_run(doc, text, size, color, bold=False, after=4, before=0, spc=None, caps=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text.upper() if caps else text)
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    if spc:
        _spc(r, spc)
    return p


def _center_rule(doc, width_in=1.8, color=None, before=6, after=8, sz="16"):
    color = A_HEX if color is None else color
    side = max(0.0, (CONTENT_W - width_in) / 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.left_indent = Inches(side)
    p.paragraph_format.right_indent = Inches(side)
    pbdr = _el("w:pBdr")
    pbdr.append(_el("w:bottom", **{"w:val": "single", "w:sz": sz, "w:space": "1", "w:color": color}))
    p._p.get_or_add_pPr().append(pbdr)
    return p


def mou_title_page(doc):
    """Cover page: title, project, party names, and date only."""
    for _ in range(3):
        doc.add_paragraph()
    _center_run(doc, "WHMIS  ·  PHARMACEUTICAL DISTRIBUTION ERP", 10, A_MAIN,
                bold=True, spc=2.0, after=2)
    _center_rule(doc, width_in=1.4, before=4, after=14, sz="18")
    _center_run(doc, "Memorandum", 40, SLATE, bold=True, after=0)
    _center_run(doc, "of Understanding", 40, SLATE, bold=True, after=10)
    _center_run(doc, "Project / System: WHMIS — Pharmaceutical Distribution ERP", 11, GREY, after=0)
    _center_rule(doc, width_in=2.4, before=16, after=16, sz="12")

    _center_run(doc, "BETWEEN", 9, GREY, bold=True, spc=1.5, after=4)
    _center_run(doc, VENDOR, 17, SLATE, bold=True, after=1)
    _center_run(doc, "Development / Technology Partner", 9.5, A_MAIN, bold=True, after=8)
    _center_run(doc, "AND", 9, GREY, bold=True, spc=1.5, after=8)
    _center_run(doc, CLIENT, 17, SLATE, bold=True, after=1)
    _center_run(doc, "Client", 9.5, A_MAIN, bold=True, after=0)

    _center_rule(doc, width_in=2.4, before=18, after=14, sz="12")
    _center_run(doc, f"Date: {DOC_DATE}", 11.5, SLATE, bold=True, after=2)
    _center_run(doc, f"Document Version {VERSION}", 9.5, GREY, after=0)


def predep_title_page(doc):
    """Cover page for the Pre-Deployment doc: title, project, parties, and date."""
    for _ in range(3):
        doc.add_paragraph()
    _center_run(doc, "WHMIS  ·  PHARMACEUTICAL DISTRIBUTION ERP", 10, A_MAIN,
                bold=True, spc=2.0, after=2)
    _center_rule(doc, width_in=1.4, before=4, after=14, sz="18")
    _center_run(doc, "Pre-Deployment Requirements", 30, SLATE, bold=True, after=0)
    _center_run(doc, "& Payment Schedule", 30, SLATE, bold=True, after=10)
    _center_run(doc, "Project / System: WHMIS — Pharmaceutical Distribution ERP", 11, GREY, after=0)
    _center_rule(doc, width_in=2.4, before=16, after=16, sz="12")

    _center_run(doc, "PREPARED FOR", 9, GREY, bold=True, spc=1.5, after=4)
    _center_run(doc, CLIENT, 17, SLATE, bold=True, after=1)
    _center_run(doc, "Client", 9.5, A_MAIN, bold=True, after=8)
    _center_run(doc, "PREPARED BY", 9, GREY, bold=True, spc=1.5, after=4)
    _center_run(doc, VENDOR, 17, SLATE, bold=True, after=1)
    _center_run(doc, "Development / Technology Partner", 9.5, A_MAIN, bold=True, after=0)

    _center_rule(doc, width_in=2.4, before=18, after=14, sz="12")
    _center_run(doc, f"Date: {DOC_DATE}", 11.5, SLATE, bold=True, after=2)
    _center_run(doc, f"Version {VERSION}   ·   {CLASSIFICATION}", 9.5, GREY, after=0)


def phasewise_title_page(doc):
    """Cover page for the Phase-Wise Development & Deployment Plan."""
    for _ in range(3):
        doc.add_paragraph()
    _center_run(doc, "WHMIS  ·  PHARMACEUTICAL DISTRIBUTION ERP", 10, A_MAIN,
                bold=True, spc=2.0, after=2)
    _center_rule(doc, width_in=1.4, before=4, after=14, sz="18")
    _center_run(doc, "Phase-Wise Development", 30, SLATE, bold=True, after=0)
    _center_run(doc, "& Deployment Plan", 30, SLATE, bold=True, after=10)
    _center_run(doc, "Project / System: WHMIS — Pharmaceutical Distribution ERP", 11, GREY, after=0)
    _center_rule(doc, width_in=2.4, before=16, after=16, sz="12")

    _center_run(doc, "PREPARED FOR", 9, GREY, bold=True, spc=1.5, after=4)
    _center_run(doc, CLIENT, 17, SLATE, bold=True, after=1)
    _center_run(doc, "Client", 9.5, A_MAIN, bold=True, after=8)
    _center_run(doc, "PREPARED BY", 9, GREY, bold=True, spc=1.5, after=4)
    _center_run(doc, VENDOR, 17, SLATE, bold=True, after=1)
    _center_run(doc, "Development / Technology Partner", 9.5, A_MAIN, bold=True, after=0)

    _center_rule(doc, width_in=2.4, before=18, after=14, sz="12")
    _center_run(doc, f"Date: {DOC_DATE}", 11.5, SLATE, bold=True, after=2)
    _center_run(doc, f"Version {VERSION}   ·   {CLASSIFICATION}", 9.5, GREY, after=0)


CONTENT_W = 7.4  # set per-doc after margins


# ============================ 2-PAGE DOCUMENT ============================== #
def build_two_page():
    global CONTENT_W
    set_theme(*THEME_BLUE)
    doc = Document()
    base_styles(doc, body_size=8.5, line=1.02)

    # ------------------------------ TITLE PAGE ------------------------------ #
    set_margins(doc, top=0.85, bottom=0.6, left=0.9, right=0.9)
    CONTENT_W = 8.5 - 0.9 - 0.9
    predep_title_page(doc)

    # -------------------------- CONTENT SECTION ----------------------------- #
    content = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(doc, top=0.45, bottom=0.4, left=0.55, right=0.55, section=content)
    CONTENT_W = 8.5 - 0.55 - 0.55
    content._sectPr.append(_el("w:pgNumType", **{"w:start": "1"}))  # content starts at page 1

    title_band(doc, "Pre-Deployment Requirements & Payment Schedule",
               f"Project / System: {PROJECT}   ·   Version {VERSION}   ·   {DOC_DATE}   ·   {CLASSIFICATION}")
    parties_row(doc)

    sec(doc, 1, "Purpose",
        [("Phase 1 of the WHMIS system is complete and will now run as a secure, cloud-based, "
          "internet-accessible solution. This document defines the pre-deployment requirements, "
          "domain and hosting selection, third-party charges, the development & "
          "deployment fee, milestone payments, responsibilities, the proposed initial go-live and "
          "post-deployment support. Detailed software scope is governed by the separate ", {}),
         ("WHMIS System – Phase-Wise Development & Deployment Plan", {"italic": True}),
         ("; contractual, support and legal terms are governed by the applicable Service / Software "
          "Development / Support & Maintenance agreements.", {})])

    sec(doc, 2, "Production Deployment Approach",
        [("The system is deployed online and all technical activity is managed by " + VENDOR +
          " — server setup, application and database deployment, domain/DNS, SSL/HTTPS, security, "
          "backups, environment configuration, testing, verification and go-live support. ", {}),
         (CLIENT + " approves only three things: (1) the domain, (2) the hosting plan, and "
          "(3) the applicable third-party charges. No infrastructure-level decisions are required "
          "of the client.", {"bold": True})])

    sectitle(doc, 3, "Domain Recommendations")
    table(doc,
          ["Proposed Domain", "Availability", "First-Year Price", "Renewal Price", "Recommendation"],
          [[f"[Domain Option {i}]", "[To Be Checked]", "PKR [Insert]", "PKR [Insert]", "[Recommendation]"]
           for i in range(1, 5)],
          widths=[1.9, 1.2, 1.25, 1.15, 1.8], right_cols=(2, 3))
    sec(doc, "", "",
        [("Domain cost is a third-party registrar charge, separate from the development & deployment fee. "
          "Availability and price are verified immediately before purchase; promotional first-year "
          "prices may differ from renewal rates, and USD/PKR movement or registrar taxes may apply. "
          "Purchase proceeds after " + CLIENT + " approves one option; the domain is registered "
          "under client ownership where possible.", {})], size=8)

    sectitle(doc, 4, "Recommended Hosting Plan")
    table(doc,
          ["Hosting Plan", "Billing Period", "First-Year Cost", "Renewal Cost", "Recommendation"],
          [["[Insert Provider] — [Insert Plan]", "[Annual / Monthly]", "PKR [Insert]", "PKR [Insert]",
            "Recommended for WHMIS Production"]],
          widths=[2.0, 1.2, 1.25, 1.15, 1.7], right_cols=(2, 3))
    sec(doc, "", "",
        [("The plan is selected by " + VENDOR + " based on the application's operational needs; the "
          "client approves the plan and its cost only, not the underlying infrastructure.", {})], size=8)

    sec(doc, 5, "What " + VENDOR + " Manages",
        [("Hosting setup, domain connection, SSL/HTTPS, application and database deployment, "
          "production security and environment configuration, backup configuration and management, "
          "user/admin setup, deployment QA, production testing and go-live support. Backup "
          "architecture and server-level configuration are handled by " + VENDOR + " — the client "
          "is not required to define them.", {})])

    callout(doc, "Third-Party Vendor Charges.",
            "Domain and hosting are not included in the development and deployment fee. "
            "Third-party vendor prices are subject to change according to vendor policies, promotional "
            "pricing, renewal pricing, applicable taxes, currency exchange-rate fluctuations and other "
            "charges imposed by the respective service provider; " + VENDOR + " does not control "
            "third-party pricing or renewal rates.")

    sectitle(doc, 6, "Charges & Payment Schedule")
    table(doc,
          ["Stage", "Scope", "Amount", "Payment Trigger"],
          [
              ["Initial Deployment", "Existing WHMIS core — production deployment & go-live",
               "PKR 100,000", "Before production deployment"],
              ["Phase 1", "Sample Management & Inventory Integration (FOC samples, PKR 0 revenue)",
               "PKR [Insert]", "Before Phase 1 starts, after go-live"],
              ["Phase 2", "HR, Employees & Payroll", "PKR 60,000",
               "Before Phase 2 starts, after Phase 1 accepted"],
              ["Phase 3", "Medical Reps, Doctors & Field-Visit Management", "PKR 70,000",
               "Before Phase 3 starts, after Phase 2 accepted"],
              ["Phase 4", "Doctor Activities, Agreements & Financial Tracking", "PKR 70,000",
               "Before Phase 4 starts, after Phase 3 accepted"],
          ],
          widths=[1.1, 3.05, 1.05, 2.2], right_cols=(2,),
          total_row=["Total", "Initial deployment plus four development phases", "PKR [Insert]",
                     "Domain & hosting billed separately"])
    sec(doc, "", "",
        [("Development order prioritizes ", {}),
         ("Sample Management first", {"bold": True}),
         (" (it directly extends the existing inventory workflow), then HR (the employee foundation), "
          "then Medical Representatives & Doctor Visits (built on HR), then Doctor Activities & "
          "Financial Tracking (built on the doctor/representative structure). Initial Deployment "
          "proceeds after: approval of this document; receipt of PKR 100,000; approval of the domain "
          "and hosting plan; payment/authorization of third-party domain & hosting charges; and "
          "provision of essential business information for setup.", {})], size=8)

    sectitle(doc, 7, "Initial Deployment & Go-Live Schedule")
    table(doc,
          ["Date", "Activity"],
          [
              ["24 Aug 2026", "Commercial approval and Initial Deployment payment confirmation"],
              ["24–25 Aug 2026", "Domain and hosting purchase / provisioning"],
              ["25–28 Aug 2026", f"Production environment, application & database deployment by {VENDOR}"],
              ["28 Aug 2026", "Internal production verification"],
              ["29 Aug 2026", "Proposed production go-live"],
          ],
          widths=[1.4, 5.9])
    sec(doc, "", "",
        [("These are proposed target dates dependent on timely payment, approvals and required "
          "information. 29 August 2026 is a proposed go-live target, not an unconditional contractual "
          "deadline; if the approval/payment date moves, the timeline moves accordingly.", {})], size=8)

    two_col(doc,
            "Responsibilities — " + VENDOR,
            ["Recommend domain & hosting and provide current pricing",
             "Purchase/configure approved domain & hosting after authorization",
             "Server, application & database deployment; DNS, SSL, security",
             "Backups, production config, QA, testing & go-live support",
             "Maintain technical production configuration under support arrangement"],
            "Responsibilities — " + CLIENT,
            ["Approve one domain and the recommended hosting plan",
             "Pay domain/hosting vendor charges and the applicable milestone",
             "Provide authorized representative details & business information",
             "Review delivered functionality and provide timely UAT approval",
             "Not responsible for architecture, backups, SSL, DBA or deploy config"])

    sec(doc, 8, "Post-Deployment Support & Maintenance",
        [(VENDOR + " provides support and maintenance for the deployed system under the agreed "
          "arrangement — application support, defect correction within the supported scope, backup "
          "monitoring, environment maintenance and agreed updates. ", {}),
         ("Included:", {"bold": True}),
         (" routine support of the agreed deployed system. ", {}),
         ("Additional development:", {"bold": True}),
         (" new modules, integrations or substantial changes are separate change requests / future "
          "phases. Commercial terms are documented in the final Service Agreement / Support & "
          "Maintenance Schedule.", {})])

    sec(doc, 9, "Change Requests & Taxes",
        [("The agreed development & deployment fee covers the currently agreed scope. Substantial new requirements are treated "
          "as change requests and assessed for functional, technical, effort, timeline and commercial "
          "impact; nothing additional is included without written approval. Applicable taxes, "
          "withholding, banking charges and government levies are handled per applicable law and the "
          "final invoice/payment arrangement.", {})])

    callout(doc, "Commercial Snapshot.",
            "Project: " + PROJECT + "  ·  Client: " + CLIENT + "  ·  Partner: " + VENDOR +
            ".  Total development & deployment: PKR [Insert] (Initial Deployment 100,000 · "
            "P1 Samples [Insert] · P2 HR 60,000 · P3 Medical Reps 70,000 · P4 Doctor Activities "
            "70,000).  Domain & hosting: separate third-party cost.  Proposed initial go-live: "
            "29 August 2026, subject to prerequisites.  Support & maintenance by " + VENDOR +
            " under the agreed arrangement.")

    sec(doc, 10, "Next Steps",
        [("Approve this document → review & select domain → approve hosting plan → pay domain & "
          "hosting charges → pay Initial Deployment (PKR 100,000) → " + VENDOR + " deploys → "
          "production go-live → commence Phase 1 (Sample Management) per the Phase-Wise Plan → "
          "finalize Service Agreement and Support & Maintenance terms.", {})])

    sectitle(doc, 11, "Acknowledgement & Sign-Off")
    sec(doc, "", "",
        [("Both parties confirm their understanding of the pre-deployment requirements, third-party "
          "domain and hosting costs, the initial deployment and four phased development milestones "
          "and the proposed deployment process. Detailed scope is governed by the approved "
          "Phase-Wise Development & Deployment "
          "Plan; broader contractual and support terms by the applicable Service Agreement and "
          "Support & Maintenance arrangements.", {})], size=8, after=4)
    sign_row(doc)

    footer(doc, f"{CLASSIFICATION}   ·   {VENDOR}   ·   {PROJECT}", section=content)
    doc.save(OUT_2P)
    print(f"Wrote {OUT_2P}")


# ============================ 1-PAGE MOU =================================== #
def build_mou():
    global CONTENT_W
    set_theme(*THEME_BLUE)
    doc = Document()
    base_styles(doc, body_size=10, line=1.08)

    # ------------------------------ TITLE PAGE ------------------------------ #
    set_margins(doc, top=0.85, bottom=0.6, left=0.9, right=0.9)
    CONTENT_W = 8.5 - 0.9 - 0.9
    mou_title_page(doc)

    # -------------------------- CONTENT SECTION ----------------------------- #
    content = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(doc, top=0.6, bottom=0.5, left=0.75, right=0.75, section=content)
    CONTENT_W = 8.5 - 0.75 - 0.75
    content._sectPr.append(_el("w:pgNumType", **{"w:start": "1"}))  # content starts at page 1

    title_band(doc, "Memorandum of Understanding",
               f"Project: {PROJECT} — Pharmaceutical Distribution ERP   ·   Version {VERSION}   ·   {DOC_DATE}")

    def para(parts, after=5, before=0, size=10, indent=0.0, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(before)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        for text, opts in parts:
            rr = p.add_run(text)
            rr.font.size = Pt(opts.get("size", size))
            rr.bold = opts.get("bold", False)
            rr.italic = opts.get("italic", italic)
            if "color" in opts:
                rr.font.color.rgb = opts["color"]
        return p

    def clause(num, title, parts, after=5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(6)
        chip_run(p, num, size=10)
        r = p.add_run(title + ". ")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = SLATE
        for text, opts in parts:
            rr = p.add_run(text)
            rr.font.size = Pt(opts.get("size", 10))
            rr.bold = opts.get("bold", False)
            rr.italic = opts.get("italic", False)
            if "color" in opts:
                rr.font.color.rgb = opts["color"]

    # Recital / opening.
    para([("This Memorandum of Understanding (this “MOU”) is made and entered into on "
           "______ ", {}),
          (DOC_DATE, {}),
          (" by and between:", {})], after=6)
    parties_row(doc, size=9, after=5)
    para([(VENDOR + " (the “Technology Partner”) and " + CLIENT + " (the “Client”) are "
           "each referred to as a “Party” and collectively as the “Parties”.", {})], after=6)

    clause(1, "Background & Purpose",
           [("The Technology Partner has developed Phase 1 of the " + PROJECT + " system for the Client, "
             "who wishes to operate it as a secure, cloud-based, internet-accessible solution and to "
             "proceed with its phased development. This MOU records the Parties’ mutual understanding "
             "of the pre-deployment requirements, the phased scope of work, the associated payments and "
             "the production deployment approach.", {})])

    # Clause 2 — scope + phases table.
    ph = doc.add_paragraph()
    ph.paragraph_format.space_after = Pt(3)
    ph.paragraph_format.space_before = Pt(6)
    chip_run(ph, 2, size=10)
    r = ph.add_run("Scope of Engagement — Phases, Payment Schedule & Timeline. ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = SLATE
    r = ph.add_run("The Parties intend a four-phase development and deployment engagement with a total "
                   "value of ")
    r.font.size = Pt(10)
    r = ph.add_run("PKR 300,000")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = A_DK
    r = ph.add_run(", as set out below:")
    r.font.size = Pt(10)
    table(doc,
          ["Phase", "Scope", "Amount", "Payment", "Timeline"],
          [
              ["Phase 1", "Production deployment of the existing WHMIS core", "PKR 100,000",
               "Before deployment", "24–29 Aug 2026 · go-live 29 Aug"],
              ["Phase 2", "HR, Employees & Payroll module", "PKR 60,000",
               "Before Phase 2 start", "≈ 3–4 weeks (after Phase 1)"],
              ["Phase 3", "Medical Reps, Doctors & Field Visit management", "PKR 70,000",
               "Before Phase 3 start", "≈ 3–4 weeks (after Phase 2)"],
              ["Phase 4", "Doctor activities, financial records, samples & integration", "PKR 70,000",
               "Before Phase 4 start", "≈ 3–4 weeks (after Phase 3)"],
          ],
          widths=[0.6, 2.45, 0.95, 1.2, 1.8], right_cols=(2,),
          total_row=["Total", "Full development & deployment scope", "PKR 300,000",
                     "Milestone-based", "Per Phase-Wise Plan"],
          size=8.5, header_fill=A_HEX)
    para([("Timelines for Phases 2–4 are indicative; each phase begins after acceptance of the preceding "
           "phase and is confirmed in the Phase-Wise Development & Deployment Plan.",
           {"italic": True, "size": 8, "color": GREY})], after=6)

    clause(3, "Payment & Third-Party Costs",
           [("Payment is milestone-based per the schedule above; each phase is payable before it "
             "commences, and Phase 1 is payable before production deployment. ", {}),
            ("Domain and hosting are separate third-party costs and are not included in the "
             "PKR 300,000 fee.", {"bold": True}),
            (" Applicable taxes, withholding and banking charges are handled in accordance with "
             "applicable law and the final invoice/payment arrangement.", {})])

    para([("4.  ", {"bold": True, "color": SLATE}),
          ("Deployment & Support Responsibilities. ", {"bold": True, "color": SLATE}),
          ("The Technology Partner manages the technical deployment and provides post-deployment "
           "support under the agreed arrangement; the Client approves the commercial elements and "
           "provides the required business information.", {})], after=3, before=6)
    two_col(doc,
            "Technology Partner (" + VENDOR + ")",
            ["Recommend domain & hosting and provide current pricing",
             "Purchase/configure approved domain & hosting after authorization",
             "Server, application & database deployment; DNS, SSL, security",
             "Backups, production configuration, QA, testing & go-live support",
             "Post-deployment support & maintenance under the agreed arrangement"],
            "Client (" + CLIENT + ")",
            ["Approve one domain and the recommended hosting plan",
             "Pay domain/hosting vendor charges and the applicable milestone",
             "Provide authorized representative details & business information",
             "Review delivered functionality and provide timely UAT approval",
             "Not responsible for technical architecture, backups, SSL or DBA"],
            size=8.5, after=6)

    clause(5, "Confidentiality",
           [("Each Party will treat the commercial, technical and project information exchanged under "
             "this engagement as confidential, will use it solely for the purpose of the engagement, and "
             "will not disclose it to any third party without the other Party’s prior written consent. "
             "This obligation survives the completion or termination of this MOU.", {})])

    clause(6, "Status of this MOU",
           [("This MOU records the Parties’ mutual intent and is ", {}),
            ("not the final or legally binding contract", {"bold": True}),
            (" (save for the confidentiality clause). The detailed software scope is governed by the ", {}),
            ("WHMIS System – Phase-Wise Development & Deployment Plan", {"italic": True}),
            (" and the ", {}),
            ("Pre-Deployment Requirements & Payment Schedule", {"italic": True}),
            ("; binding contractual, support and legal terms are governed by the applicable Service / "
             "Software Development / Support & Maintenance agreements to be executed by the Parties.", {})])

    clause(7, "Term & Next Steps",
           [("This MOU is effective from the date of the last signature below and remains in effect until "
             "superseded by the definitive agreements. The Parties intend to proceed by: approving this "
             "MOU; selecting and approving the domain and hosting; paying the Phase 1 milestone; "
             "completing production deployment and go-live; and commencing subsequent phases in "
             "accordance with the Phase-Wise Development & Deployment Plan.", {})])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("IN WITNESS WHEREOF, the Parties have signed this Memorandum of Understanding as of "
                  "the date first written above.")
    r.font.size = Pt(9.5)
    r.italic = True
    r.font.color.rgb = INK
    blank(doc, 4)
    sign_row(doc, size=9)

    footer(doc, f"{CLASSIFICATION}   ·   {VENDOR}   ·   {PROJECT} — Memorandum of Understanding",
           section=content)
    doc.save(OUT_MOU)
    print(f"Wrote {OUT_MOU}")


# =================== PHASE-WISE DEVELOPMENT & DEPLOYMENT PLAN ============== #
def build_phasewise():
    global CONTENT_W
    set_theme(*THEME_BLUE)
    doc = Document()
    base_styles(doc, body_size=8.5, line=1.03)

    # ------------------------------ TITLE PAGE ------------------------------ #
    set_margins(doc, top=0.85, bottom=0.6, left=0.9, right=0.9)
    CONTENT_W = 8.5 - 0.9 - 0.9
    phasewise_title_page(doc)

    # -------------------------- CONTENT SECTION ----------------------------- #
    content = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(doc, top=0.45, bottom=0.4, left=0.55, right=0.55, section=content)
    CONTENT_W = 8.5 - 0.55 - 0.55
    content._sectPr.append(_el("w:pgNumType", **{"w:start": "1"}))

    title_band(doc, "Phase-Wise Development & Deployment Plan",
               f"Project / System: {PROJECT}   ·   Version {VERSION}   ·   {DOC_DATE}   ·   {CLASSIFICATION}")

    def objective(text):
        sec(doc, "", "", [("Objective. ", {"bold": True, "color": A_DK}), (text, {})],
            size=8, after=2)

    def phase_block(num, title, obj, left_title, left_items, right_title, right_items):
        sectitle(doc, num, title)
        objective(obj)
        two_col(doc, left_title, left_items, right_title, right_items, size=8, after=4)

    # 1 — Purpose & Scope
    sec(doc, 1, "Purpose & Scope",
        [("This plan defines the detailed functional scope, delivery sequence and acceptance basis "
          "for the " + PROJECT + " engagement between " + VENDOR + " and " + CLIENT + ". It governs the "
          "software scope referenced by the ", {}),
         ("Pre-Deployment Requirements & Payment Schedule", {"italic": True}),
         (" and the ", {}),
         ("Memorandum of Understanding", {"italic": True}),
         (". Delivery comprises four phases: Phase 1 deploys the existing WHMIS core to production "
          "together with Sample Management & Inventory Integration; three further development phases "
          "follow — each demonstrated, accepted and paid before the next begins.", {})])

    # 2 — Delivery Model & Sequence
    sec(doc, 2, "Delivery Model & Sequence",
        [("Work is delivered incrementally on the live production system, each phase building on the "
          "one before it. ", {}),
         ("Phase 1", {"bold": True}),
         (" deploys the existing WHMIS core to production together with ", {}),
         ("Sample Management & Inventory Integration", {"bold": True}),
         (", which directly extends the existing Purchase, Inventory and Sales workflows. ", {}),
         ("Phase 2 — HR, Employees & Payroll", {"bold": True}),
         (" creates the employee foundation; ", {}),
         ("Phase 3 — Medical Representatives & Doctor Visits", {"bold": True}),
         (" builds on that HR foundation; and ", {}),
         ("Phase 4 — Doctor Activities & Financial Tracking", {"bold": True}),
         (" builds on the resulting doctor and representative structure.", {})])

    # 3 — Scope & Payment Summary
    sectitle(doc, 3, "Scope, Sequence & Payment Summary")
    table(doc,
          ["Stage", "Scope Summary", "Amount", "Precedence"],
          [
              ["Phase 1", "Existing WHMIS core + Sample Management & Inventory Integration "
               "to production & go-live", "PKR [Insert]", "First — on approval & prerequisites"],
              ["Phase 2", "HR, Employees & Payroll", "PKR 60,000",
               "After Phase 1 go-live"],
              ["Phase 3", "Medical Reps, Doctors & Field-Visit Management", "PKR 70,000",
               "After Phase 2 accepted"],
              ["Phase 4", "Doctor Activities, Agreements & Financial Tracking", "PKR 70,000",
               "After Phase 3 accepted"],
          ],
          widths=[1.1, 3.0, 1.05, 2.25], right_cols=(2,),
          total_row=["Total", "All four development & deployment phases", "PKR [Insert]",
                     "Milestone-based"])
    sec(doc, "", "",
        [("Phase 1 covers the existing WHMIS core (PKR 100,000) plus Sample Management & Inventory "
          "Integration; the combined amount is confirmed on approval. Domain and hosting are separate "
          "third-party charges.", {})], size=8)

    # 4 — Initial Deployment (existing core + Sample Management)
    sectitle(doc, 4, "Phase 1 — Existing WHMIS Core & Sample Management")
    objective(
        "Deploy the already-developed WHMIS core to a secure, internet-accessible production "
        "environment and, as part of the same Phase 1 delivery, add Sample Management & Inventory "
        "Integration, then achieve a verified go-live.")
    two_col(doc,
            "Deployment Activities",
            ["Provision hosting, connect the domain and enable SSL/HTTPS",
             "Deploy the application and database to production",
             "Configure the production environment, security and access",
             "Configure and schedule automated backups"],
            "Go-Live & Verification",
            ["Create admin/user accounts and roles",
             "Deployment QA and production smoke testing",
             "Setup and data verification with the client",
             "Production go-live and handover support"],
            size=8, after=3)
    sec(doc, "", "",
        [("Included — Sample Management & Inventory Integration. ", {"bold": True, "color": A_DK}),
         ("Free-of-cost (FOC) samples are managed end to end, extending the existing Purchase, "
          "Inventory and Sales workflows without disturbing normal stock or revenue.", {})],
        size=8, after=2)
    two_col(doc,
            "Sample Capabilities",
            ["Track sample quantities separately from saleable stock",
             "Receive samples through purchasing / stock receipt",
             "Maintain batch and expiry information where applicable",
             "Issue FOC samples to doctors/recipients; reduce inventory on issue",
             "Link sample issuance to the responsible Medical Representative where required",
             "Retain inventory cost while sample sale/revenue value stays PKR 0"],
            "Sample Documents & Reports",
            ["Sample / FOC Issue document generation",
             "Sample stock report",
             "Sample movement report",
             "Product-wise sample report",
             "Doctor-wise sample report",
             "Representative-wise sample report"],
            size=8, after=4)

    # 5 — Phase 2 HR
    phase_block(
        5, "Phase 2 — HR, Employees & Payroll",
        "Establish the employee foundation with master records, structured salaries and "
        "salary-slip generation.",
        "Employees & Structure",
        ["Employee master records",
         "Departments and designations",
         "Employment information and employee status",
         "Link application users with employee profiles",
         "Medical Representatives maintained as employees",
         "Salary structures with allowances and deductions"],
        "Payroll & Reporting",
        ["Monthly payroll processing",
         "Salary history per employee",
         "Salary slip generation",
         "Printable / PDF salary slips",
         "HR and payroll reporting"])

    # 6 — Phase 3 MR / Doctors / Field Visits
    phase_block(
        6, "Phase 3 — Medical Representatives, Doctors & Field Visit Management",
        "Manage Medical Representatives, the doctor directory and field-visit activity with "
        "supporting evidence and management monitoring.",
        "Setup & Field Visits",
        ["MR profiles linked to employees/users",
         "Doctor master directory",
         "Doctor assignment to MRs; territory/area assignment",
         "Visit logging: date, time, purpose, notes, outcome",
         "Follow-up requirements and next visit dates",
         "Photo and document attachments as visit evidence"],
        "History, Monitoring & Reports",
        ["Doctor-wise visit history",
         "MR-wise visit history",
         "Management monitoring dashboard",
         "Visit reports by date and doctor",
         "Visit reports by representative and territory",
         "Visit reports by status"])

    # 7 — Phase 4 Doctor Activities & Financial Tracking
    phase_block(
        7, "Phase 4 — Doctor Activities, Agreements & Financial Tracking",
        "Record doctor activities and arrangements and track their committed, paid and outstanding "
        "amounts against expected and actual returns.",
        "Activities & Agreements",
        ["Doctor activities / arrangements",
         "Discount / percentage-based arrangements",
         "Bonus arrangements",
         "Promotional / activity-based arrangements",
         "Agreed amounts and percentages",
         "MR association with doctor activities"],
        "Financial Tracking & Reports",
        ["Expected vs. actual business/return tracking",
         "Payment logs with multiple payments per activity",
         "Committed / paid / outstanding amounts",
         "Supporting documents / attachments",
         "Doctor-wise activity & financial ledger",
         "Expected-vs-actual return reporting & dashboards"])

    # 8 — Acceptance, UAT & Timeline
    sec(doc, 8, "Acceptance, UAT & Timeline",
        [("Each phase concludes with client user-acceptance testing (UAT) and written acceptance; the "
          "next phase begins only after that acceptance and its milestone payment. Phase 1 "
          "targets a proposed production go-live of 29 August 2026, subject to timely approval, payment "
          "and provision of required information. Individual development phases are indicatively ", {}),
         ("3–4 weeks", {"bold": True}),
         (" each and are confirmed at the start of each phase; dates move if approvals or payments move.", {})])

    # 9 — Assumptions, Exclusions & Change Control
    sectitle(doc, 9, "Assumptions, Exclusions & Change Control")
    two_col(doc,
            "Assumptions",
            ["Timely payments, approvals and UAT sign-offs",
             "Client provides required business data and authorized contacts",
             "Built on the existing WHMIS platform and its current data model",
             "One production environment; standard operating practices"],
            "Exclusions & Change Control",
            ["Domain, hosting and third-party charges are billed separately",
             "New modules/integrations beyond this scope are change requests",
             "Substantial changes assessed for functional, technical, effort, timeline & commercial impact",
             "Nothing additional is included without written approval"],
            size=8, after=4)

    # 10 — Acknowledgement & Sign-Off
    sectitle(doc, 10, "Acknowledgement & Sign-Off")
    sec(doc, "", "",
        [("Both parties confirm their understanding of the phased scope, the delivery sequence and the "
          "acceptance basis set out above. This plan governs the software scope only; commercial terms "
          "are set out in the Pre-Deployment Requirements & Payment Schedule, and binding contractual, "
          "support and legal terms in the applicable Service / Support & Maintenance agreements.", {})],
        size=8, after=4)
    sign_row(doc)

    footer(doc, f"{CLASSIFICATION}   ·   {VENDOR}   ·   {PROJECT} — Phase-Wise Plan", section=content)
    doc.save(OUT_PW)
    print(f"Wrote {OUT_PW}")


# ================================ INVOICE ================================= #
def build_invoice():
    global CONTENT_W
    set_theme(*THEME_BLUE)
    doc = Document()
    set_margins(doc, top=0.5, bottom=0.45, left=0.5, right=0.5)
    CONTENT_W = 8.5 - 0.5 - 0.5
    base_styles(doc, body_size=9, line=1.05)

    INV_NO = "VWT-INV-2026-0001"
    INV_ISSUE = "21 August 2026"
    INV_DUE = "Before production deployment"

    # ------------------------------- HEADER --------------------------------- #
    hdr = doc.add_table(rows=1, cols=2)
    hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr.autofit = False
    L, R = hdr.rows[0].cells
    L.width = Inches(4.4)
    R.width = Inches(CONTENT_W - 4.4)

    _cell_margins(L, top=60, bottom=60, left=10, right=90)
    L.text = ""
    p = L.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("WHMIS  ·  PHARMACEUTICAL DISTRIBUTION ERP")
    r.bold = True
    r.font.size = Pt(7.5)
    r.font.color.rgb = A_MAIN
    _spc(r, 1.0)
    p = L.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(VENDOR)
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = SLATE
    p = L.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Development / Technology Partner")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = A_MAIN
    for line in ("pm@vwisdomtechnologies.com   ·   +92 318 5161571",
                 "Contact: Muhammad Saad Mubeen, Project Manager"):
        p = L.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        r.font.size = Pt(8)
        r.font.color.rgb = GREY

    _shade(R, A_HEX)
    _cell_margins(R, top=90, bottom=90, left=120, right=120)
    R.text = ""
    p = R.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("INVOICE")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = WHITE
    _spc(r, 2.0)
    for label, value in (("Invoice No", INV_NO), ("Issue Date", INV_ISSUE),
                         ("Payment Due", INV_DUE), ("Currency", "PKR (Pakistani Rupee)")):
        pl = R.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pl.paragraph_format.space_after = Pt(1)
        rl = pl.add_run(label + ":  ")
        rl.font.size = Pt(8)
        rl.font.color.rgb = BAND_SUB
        rv = pl.add_run(value)
        rv.bold = True
        rv.font.size = Pt(8.5)
        rv.font.color.rgb = WHITE
    blank(doc, 6)

    # ------------------------------ BILL TO --------------------------------- #
    bt = doc.add_table(rows=1, cols=2)
    bt.alignment = WD_TABLE_ALIGNMENT.CENTER
    bt.autofit = False
    half = CONTENT_W / 2
    billed = [
        ("BILL TO", CLIENT,
         [("Attn: ", "Mr. Jamil Ahmed"), ("Address: ", "[Insert]"),
          ("NTN / STRN: ", "[Insert]")]),
        ("PROJECT / REFERENCE", "WHMIS — Pharmaceutical Distribution ERP",
         [("Scope: ", "Initial Deployment + Phases 1–4"),
          ("Governed by: ", "Pre-Deployment Schedule & Phase-Wise Plan"),
          ("Reference: ", INV_NO)]),
    ]
    for cell, (label, name, lines) in zip(bt.rows[0].cells, billed):
        cell.width = Inches(half)
        _cell_margins(cell, top=80, bottom=75, left=120, right=110)
        _cell_border(cell, "top", 22, A_HEX)
        for edge in ("bottom", "left", "right"):
            _cell_border(cell, edge, 4, RULE)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(7)
        r.font.color.rgb = A_MAIN
        _spc(r, 0.8)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        r = p2.add_run(name)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = SLATE
        for lab, val in lines:
            pl = cell.add_paragraph()
            pl.paragraph_format.space_after = Pt(1)
            rl = pl.add_run(lab)
            rl.font.size = Pt(8)
            rl.font.color.rgb = GREY
            rv = pl.add_run(val)
            rv.font.size = Pt(8)
            rv.font.color.rgb = INK
    blank(doc, 6)

    # -------------------------- DUE NOW LINE ITEMS -------------------------- #
    sectitle(doc, 1, "Amount Due Now — Payable Before Deployment")
    table(doc,
          ["Description", "Type", "Amount"],
          [
              ["Domain registration — [Domain Option] (1 year)", "Third-party pass-through",
               "PKR [Insert]"],
              ["Production hosting plan (annual)", "Third-party pass-through", "PKR [Insert]"],
              ["Phase 1 (Initial Deployment) — existing WHMIS core to production & go-live · "
               "Installment 1 of 2",
               "Development & deployment", "PKR 50,000"],
          ],
          widths=[4.15, 2.25, 1.1], right_cols=(2,), size=8.5,
          total_row=["Total Due Now", "Domain + hosting + Phase 1 first installment", "PKR [Insert]"])
    callout(doc, "Amount Due Now:",
            "PKR [Insert]  (domain + hosting + Phase 1 first installment of PKR 50,000).  Payable "
            "before production deployment begins; the remaining PKR 50,000 of Phase 1 falls due one "
            "month later (see below). Domain and hosting are third-party pass-through charges billed "
            "at actual provider cost; " + VENDOR + " does not control third-party pricing, renewals, "
            "taxes or currency movement.", size=8.5)

    # ---------------------- SCHEDULED FUTURE PAYMENTS ----------------------- #
    sectitle(doc, 2, "Scheduled Future Payments — Per Phase Deployment")
    table(doc,
          ["Milestone", "Scope", "Expected Due Date", "Amount"],
          [
              ["Phase 1", "Existing WHMIS core (incl. Sample Management, FOC) · Installment 2 of 2",
               "Est. 21 Sep 2026 · one month after go-live", "PKR 50,000"],
              ["Phase 2", "HR, Employees & Payroll",
               "Est. 29 Sep 2026 · after Phase 1 accepted", "PKR 60,000"],
              ["Phase 3", "Medical Reps, Doctors & Field Visits",
               "Est. 27 Oct 2026 · after Phase 2 accepted", "PKR 70,000"],
              ["Phase 4", "Doctor Activities, Agreements & Financial Tracking",
               "Est. 24 Nov 2026 · after Phase 3 accepted", "PKR 70,000"],
          ],
          widths=[0.85, 2.6, 2.55, 1.0], right_cols=(3,), size=8.5,
          total_row=["", "Future payments subtotal (Phase 1 balance + Phases 2–4)", "", "PKR 250,000"])
    sec(doc, "", "",
        [("Each future milestone is invoiced separately when its phase begins. Expected due dates are "
          "indicative and move with approvals, payments and phase acceptance; each phase is payable "
          "before it commences.", {})], size=8)

    callout(doc, "Contract Summary.",
            "Total development & deployment: PKR 300,000 (Phase 1 100,000 — payable as 50,000 now + "
            "50,000 one month later · P2 HR 60,000 · P3 Medical Reps 70,000 · P4 Doctor Activities "
            "70,000).  Domain & hosting are separate third-party costs, billed at actual provider "
            "cost.", size=8.5)

    # --------------------- REMIT-TO + NOTES (two columns) ------------------- #
    two_col(doc,
            "Remit To",
            ["Account Title: [Insert]",
             "Bank / Branch: [Insert]",
             "Account No / IBAN: [Insert]",
             "Payment Reference: " + INV_NO],
            "Notes",
            ["Future milestones are invoiced separately at each phase start",
             "Expected due dates are indicative; subject to acceptance",
             "Domain & hosting: third-party pricing, taxes & FX may apply",
             "Applicable taxes, withholding & bank charges per law"],
            size=8, after=5)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Thank you for your business.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = A_DK

    footer(doc, f"{CLASSIFICATION}   ·   {VENDOR}   ·   {PROJECT} — Invoice {INV_NO}")
    doc.save(OUT_INV)
    print(f"Wrote {OUT_INV}")


if __name__ == "__main__":
    os.makedirs(OUTDIR, exist_ok=True)
    build_two_page()
    build_mou()
    build_phasewise()
    build_invoice()
