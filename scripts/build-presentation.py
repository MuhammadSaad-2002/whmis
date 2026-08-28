#!/usr/bin/env python3
"""Build the WHMIS client-facing system presentation as a formatted .docx.

Consumes the screenshots in docs/presentation/img/ (captured by
scripts/capture-screenshots.mjs) and emits
docs/presentation/WHMIS-System-Presentation.docx.

    python3 scripts/build-presentation.py
"""
import os
from datetime import date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMG = os.path.join(ROOT, "docs", "presentation", "img")
OUT = os.path.join(ROOT, "docs", "presentation", "WHMIS-System-Presentation.docx")

# Identity
CLIENT = "MASTER PHARMACUTICALS DISTRIBUTOR"
VENDOR = "Virtual Wisdom Technologies (SMC-Private) Limited"
YEAR = date.today().year

# Palette (matches the app's slate + emerald look).
SLATE = RGBColor(0x0F, 0x17, 0x2A)
EMERALD = RGBColor(0x05, 0x96, 0x69)
EMERALD_DK = RGBColor(0x04, 0x78, 0x57)
GREY = RGBColor(0x6B, 0x72, 0x80)
LIGHT = RGBColor(0x9C, 0xA3, 0xAF)
CALLOUT_FILL = "ECFDF5"   # emerald-50
RULE = "D1D5DB"           # grey-300
CONTENT_WIDTH = 6.5       # inches, for 1-inch margins on Letter

# --------------------------------------------------------------------------- #
# Content model
# --------------------------------------------------------------------------- #
EXEC = {
    "paras": [
        "WHMIS is a complete, purpose-built distribution management platform for pharmaceutical "
        "distributors. It unifies every step of the business — purchasing stock and tracking "
        "batches, booking orders from the field, invoicing pharmacies, collecting payments, and "
        "analysing profit — inside a single, fast, secure web application. Instead of juggling "
        "spreadsheets, hand-written registers and disconnected tools, the entire team works from "
        "one accurate, always-current source of truth.",
        "Pharmaceutical distribution is uniquely demanding. Products move in batches with expiry "
        "dates; suppliers reward you with free “bonus” goods that quietly change your real cost; "
        "pharmacies buy on credit; and margins are thin. Small, everyday errors — an expiring batch "
        "shipped late, a customer pushed past their credit limit, a trade scheme applied "
        "inconsistently — erode profit silently, month after month. WHMIS is engineered around "
        "exactly these realities, so the system protects your margin while your team simply does "
        "its job.",
    ],
    "why": (
        f"Why WHMIS fits {CLIENT}",
        "Your business runs on batches, bonus stock, credit customers and trade schemes — the very "
        "things generic accounting software handles poorly. WHMIS was built for pharmaceutical "
        "distribution first, so the controls you need are already in place on day one, not bolted "
        "on later.",
    ),
    "gains": [
        "One connected workflow — purchase → stock → book → sell → collect → report — with every figure reconciled automatically.",
        "Purpose-built for pharma: batch and expiry tracking, bonus-stock handling, and a true cost that includes free goods.",
        "Localised for Pakistan: prices in PKR, with GST and NTN / STRN / CNIC fields on every party and document.",
        "Fast, keyboard-driven data entry so your team bills in seconds, not minutes.",
        "Runs on ordinary cPanel web hosting and works from any browser — no special servers or installs.",
        "Secure by role, so every staff member sees only what their job requires.",
    ],
}

SECTIONS = [
    {
        "cat": "OVERVIEW",
        "title": "Executive Dashboard — Your Business at a Glance",
        "overview":
            "The dashboard is the first screen your team sees each day, and it is designed to answer "
            "the questions an owner actually asks: How are we selling? How much are we making? Who "
            "owes us money, and who do we owe? It pulls live figures from across the system and "
            "presents them as clear cards and a twelve-month trend, so you understand the health of "
            "the business in seconds — without opening a single report or spreadsheet.",
        "why": (
            "Why this matters",
            "Most distributors only discover a bad month when the accounts are closed weeks later. "
            "A live dashboard turns that hindsight into foresight — you see receivables climbing, "
            "profit dipping or stock value ballooning while there is still time to act.",
        ),
        "capabilities": [
            ("Today at a glance", "today’s sales and purchases, plus this month’s sales and profit, shown the moment you log in."),
            ("Cash-flow visibility", "total receivable from customers and total payable to suppliers, always current."),
            ("Inventory value", "the worth of everything you hold, valued at true cost, so capital tied up in stock is never a mystery."),
            ("Work-in-progress", "counts of draft invoices and bookings awaiting approval, so nothing is left unfinished."),
            ("Twelve-month trend", "a sales-and-profit chart that makes momentum — or a slowdown — impossible to miss."),
        ],
        "gains": [
            "Decisions based on today’s numbers, not last month’s.",
            "Early warning on receivables, profit and overstock.",
            "Less time compiling figures, more time acting on them.",
        ],
        "images": [("02-dashboard.png", "The dashboard: the key numbers and a 12-month trend on one screen.")],
    },
    {
        "cat": "TRANSACTIONS",
        "title": "Sales Invoicing — Bill Accurately in Seconds",
        "overview":
            "Sales is the heart of your day, and WHMIS makes it fast and error-proof. Invoices are "
            "raised on a keyboard-driven grid that a trained operator can fly through without ever "
            "reaching for the mouse. As each line is entered, the system applies the right price, "
            "discount and GST, pulls stock from the correct batch, and checks the customer’s credit "
            "— then re-verifies every figure on the server before the invoice is posted, so a "
            "mistyped total can never slip through.",
        "why": (
            "Why this matters",
            "Every invoice touches stock, pricing, credit and profit at once. When those are handled "
            "by hand, errors are inevitable — the wrong batch shipped, an expired product sold, a "
            "risky customer over-extended. WHMIS removes that risk from the busiest, most repetitive "
            "task in your business.",
        ),
        "capabilities": [
            ("High-speed entry", "add products, quantities and bonus units entirely from the keyboard, built for high daily volumes."),
            ("Automatic pricing", "line discounts, GST and margins are computed instantly and re-checked by the server on save."),
            ("Batch-aware stock", "goods are drawn from the earliest-expiry batch automatically, so you never oversell or ship expired stock."),
            ("Credit control", "customers over their credit limit are blocked before a credit invoice can be posted."),
            ("Cash and credit", "cash invoices settle themselves on posting; credit sales flow straight to the customer’s ledger."),
            ("Professional output", "every invoice prints or shares as a clean, branded PDF."),
        ],
        "gains": [
            "Faster billing counters and shorter queues.",
            "Fewer pricing disputes and zero expired-stock shipments.",
            "Credit exposure controlled automatically, protecting cash flow.",
        ],
        "images": [
            ("03-sales-index.png", "Sales register with gross, returns, net and profit summarised at the top."),
            ("04-sales-entry.png", "The fast invoice-entry grid, built for high-volume billing."),
        ],
    },
    {
        "cat": "TRANSACTIONS",
        "title": "Bookings & Approvals — From Field Order to Invoice",
        "overview":
            "Field bookers capture pharmacy orders on the same fast grid your billing team uses, and "
            "each order flows through a clear submit-and-approve workflow. A manager reviews, then "
            "approves or rejects with a single click; an approved booking becomes a ready-to-bill "
            "draft sales invoice in one step, carrying its pricing and trade schemes with it. Nothing "
            "is re-keyed, and nothing is lost between the field and the invoice.",
        "why": (
            "Why this matters",
            "Orders taken on paper or over the phone get mis-transcribed, forgotten or billed at the "
            "wrong price. A structured booking pipeline gives you accountability — you know who "
            "booked what, who approved it, and that it will be invoiced exactly as agreed.",
        ),
        "capabilities": [
            ("Familiar order entry", "bookers use the same grid as sales, so training is minimal and errors are rare."),
            ("Approval workflow", "every booking is submitted for review and approved or rejected by a manager."),
            ("One-click conversion", "an approved booking becomes a draft sales invoice instantly, with no re-entry."),
            ("Stacked trade schemes", "multiple incentives — say a free-goods deal plus a percentage discount — can be applied to one line and carried through to the sale."),
            ("Clear ownership", "bookers see only their own customers’ orders, keeping the pipeline tidy and accountable."),
        ],
        "gains": [
            "A disciplined, auditable order-to-invoice pipeline.",
            "No lost orders and no re-keying between field and office.",
            "Trade schemes honoured exactly as promised to the pharmacy.",
        ],
        "images": [
            ("05-bookings-index.png", "The bookings pipeline, tracked by status through to conversion."),
            ("06-bookings-entry.png", "Order entry for field bookers, mirroring the sales experience."),
        ],
    },
    {
        "cat": "TRANSACTIONS",
        "title": "Purchasing & Bonus Stock — Know Your True Cost",
        "overview":
            "Purchases are recorded with the detail that matters in pharma — batch numbers, expiry "
            "dates, and the free “bonus” units suppliers give you. On posting, stock is created "
            "automatically and the free goods are folded into the batch’s effective cost. That single "
            "piece of accounting discipline means your profit figures are honest: the free units you "
            "later give away are already accounted for, so you are never fooled by margins that look "
            "better than they really are.",
        "why": (
            "Why this matters",
            "Bonus stock is where distributor profit quietly leaks. If free goods are ignored in "
            "costing, every report overstates your margin — and you make pricing and scheme decisions "
            "on numbers that are wrong. WHMIS gets this right by design.",
        ),
        "capabilities": [
            ("Batch & expiry capture", "record batch number and expiry on every purchase line; stock is created automatically on posting."),
            ("True cost with bonus", "free units dilute the batch’s effective cost, so profit is always calculated truthfully."),
            ("Supplier ledger", "credit and cash purchases update supplier payables and the ledger automatically."),
            ("Controlled corrections", "posted documents are never silently edited; changes are made through auditable reversals."),
        ],
        "gains": [
            "Profit reporting you can actually trust.",
            "Complete batch and expiry traceability from the moment stock arrives.",
            "Supplier balances that are always accurate.",
        ],
        "images": [
            ("07-purchases-index.png", "Purchase register with per-supplier volume and status."),
            ("08-purchases-entry.png", "Purchase entry capturing batch, expiry and bonus units."),
        ],
    },
    {
        "cat": "INVENTORY",
        "title": "Inventory, Batches & Expiry — Stop Losing Stock to Waste",
        "overview":
            "WHMIS tracks stock not just by product but by individual batch, each with its own expiry "
            "date and true cost. When goods are sold, the system consumes the earliest-expiring batch "
            "first, actively pushing older stock out before it can expire. Near-expiry and low-stock "
            "situations are surfaced early, and every movement in and out is recorded in an "
            "append-only history you can always trust.",
        "why": (
            "Why this matters",
            "Expired stock is money thrown away, and stock-outs are lost sales. For a pharmaceutical "
            "distributor, expiry management is not a nicety — it is a direct line to your bottom line. "
            "WHMIS turns expiry control from a manual worry into an automatic behaviour.",
        ),
        "capabilities": [
            ("Live batch stock", "current quantities by product and by batch, valued at true cost."),
            ("Expiry-first issue (FIFO)", "the earliest-expiring batch is always consumed first, minimising write-offs."),
            ("Early warnings", "near-expiry and low-stock visibility so you act before a loss or a shortage happens."),
            ("Full traceability", "an append-only movement history records every unit’s journey in and out."),
        ],
        "gains": [
            "Fewer expiry write-offs and fewer stock-outs.",
            "Confidence in exactly what you hold and what it is worth.",
            "A clean audit trail for every unit of stock.",
        ],
        "images": [
            ("09-inventory.png", "Current stock and value across your catalogue."),
            ("10-batches.png", "Batch-level detail with expiry dates and quantities."),
        ],
    },
    {
        "cat": "MASTER DATA",
        "title": "Products & Suppliers — A Clean Foundation",
        "overview":
            "Everything the system does rests on well-organised master data, and WHMIS makes it easy "
            "to set up and maintain. Products carry trade and retail pricing, tax and supplier links, "
            "and can be grouped by category; suppliers hold the full tax details needed for compliant "
            "documents. When you are starting out or onboarding a new range, you can import products "
            "and customers in bulk from Excel, with each row validated and clear errors shown for "
            "anything that needs attention.",
        "why": (
            "Why this matters",
            "Messy product and price lists cause wrong invoices, wrong margins and wasted time. "
            "Getting master data right once — and being able to bulk-load it — saves countless hours "
            "and prevents errors everywhere else in the system.",
        ),
        "capabilities": [
            ("Complete product records", "trade and retail prices, tax, supplier and category on every product."),
            ("Compliant supplier records", "full NTN / STRN and contact details for correct documentation."),
            ("Bulk Excel import", "load products and customers in bulk, with per-row validation and downloadable templates."),
            ("Simple organisation", "categories keep large catalogues easy to navigate."),
        ],
        "gains": [
            "A fast, error-free start with bulk onboarding.",
            "Correct pricing and tax on every document that follows.",
            "A tidy catalogue that scales as your range grows.",
        ],
        "images": [
            ("11-products.png", "The product catalogue with pricing and supplier links."),
            ("12-suppliers.png", "Supplier directory with tax and contact details."),
        ],
    },
    {
        "cat": "FINANCE",
        "title": "Customers, Ledger & Aging — Know Who Owes You",
        "overview":
            "Every customer has a complete profile — credit limit, assigned booker and tax details — "
            "and a running ledger that captures each invoice, receipt and adjustment in one place. "
            "Outstanding balances are automatically bucketed by age, so your recovery team can see at "
            "a glance which accounts are current, which are stretching, and which are genuinely "
            "overdue, and chase the right money first.",
        "why": (
            "Why this matters",
            "In a credit-driven business, cash is only as healthy as your collections. Without clear, "
            "aged balances, overdue money hides in plain sight and working capital quietly drains "
            "away. Aging turns a vague worry into a prioritised action list.",
        ),
        "capabilities": [
            ("Rich customer profiles", "credit limits, assigned bookers and tax fields on every pharmacy."),
            ("Per-customer ledger", "every invoice, receipt and adjustment in a single, clear statement."),
            ("Aged outstanding", "balances split into age buckets so overdue accounts stand out immediately."),
            ("Credit awareness", "limits feed directly into sales, blocking risky exposure at the point of billing."),
        ],
        "gains": [
            "Faster, smarter collections and healthier cash flow.",
            "Instant answers to “how much does this pharmacy owe us?”",
            "Credit risk controlled before it becomes bad debt.",
        ],
        "images": [
            ("13-customers.png", "Customer directory with credit limits and balances."),
            ("14-customer-ledger.png", "A customer’s running ledger and outstanding position."),
        ],
    },
    {
        "cat": "SALES ENABLEMENT",
        "title": "Incentive Rules — Your Trade Schemes, Applied Consistently",
        "overview":
            "Trade schemes drive sales in pharma, but applied by hand they are inconsistent and hard "
            "to measure. WHMIS lets you define each scheme once — free-goods bonuses like 10 + 2, "
            "slab-based bonuses, percentage or fixed discounts, or special prices — and target it to "
            "a specific customer, product or supplier, or make it global, with clear priorities. The "
            "rule simply fills in the invoice line, so your team can still adjust manually, and an "
            "“Incentives Given” report shows exactly what every scheme has cost you.",
        "why": (
            "Why this matters",
            "Schemes are a major cost of doing business, yet most distributors cannot say what they "
            "actually spend on them. Defining schemes centrally means every customer is treated "
            "consistently and fairly — and, for the first time, you can measure the return on that "
            "spend.",
        ),
        "capabilities": [
            ("Every scheme type", "quantity bonus, slab bonus, percentage and fixed discounts, and special prices."),
            ("Precise targeting", "apply a scheme to a customer, product or supplier — or globally — with priorities to resolve overlaps."),
            ("Manual override", "rules only fill the line, so your staff always keep the final say."),
            ("Full measurement", "the Incentives Given report reveals the true cost of every scheme, per customer."),
        ],
        "gains": [
            "Consistent, fair pricing across every pharmacy.",
            "Trade-scheme spend you can finally see and control.",
            "Faster billing, because schemes apply themselves.",
        ],
        "images": [("15-incentives.png", "The incentive-rule library, covering every scheme type.")],
    },
    {
        "cat": "FINANCE",
        "title": "Payments & Receipts — Keep Every Balance Current",
        "overview":
            "Customer receipts and supplier payments are recorded across cash and bank and allocated "
            "against the invoices they settle. Balances and aging update the instant a payment is "
            "entered, cash sales settle themselves on posting, and clear running totals show money "
            "in, money out and your net position — so your books are always up to date, not "
            "reconciled at month-end.",
        "why": (
            "Why this matters",
            "When payments are logged separately from invoices, balances drift and disputes follow. "
            "Allocating receipts against invoices as they happen keeps every customer and supplier "
            "statement accurate and defensible at all times.",
        ),
        "capabilities": [
            ("Receipts and payments", "record both, across cash and bank, in one place."),
            ("Invoice allocation", "apply money received to the specific invoices it clears."),
            ("Instant balances", "aging and outstanding update immediately, with cash sales auto-settled."),
            ("Net position", "running totals of inflow, outflow and the balance between them."),
        ],
        "gains": [
            "Always-accurate customer and supplier statements.",
            "Fewer payment disputes and faster reconciliation.",
            "A live view of cash movement.",
        ],
        "images": [("16-payments.png", "Receipts and payments with a live net position.")],
    },
    {
        "cat": "FINANCE",
        "title": "Financial Position — What You Own vs. What You Owe",
        "overview":
            "The Financial Position screen answers the single most important question an owner asks: "
            "where do we stand right now? On one page it sets what your customers owe you against "
            "what you owe your suppliers and shows the net figure between them. Three headline cards — "
            "Due from Customers, Owed to Suppliers and Net Position — sit above a full receivables "
            "statement (with aging buckets and how much each pharmacy has paid in the period) and a "
            "matching payables statement, followed by a complete log of every payment in and out for "
            "the chosen dates. Pick any date range and print the whole statement to PDF.",
        "why": (
            "Why this matters",
            "Profit on paper means little if the cash is stuck in receivables while supplier bills "
            "come due. A single, current view of receivables versus payables — and the net position "
            "between them — is what tells you whether the business can meet its obligations this "
            "week. Without it, that judgement is guesswork.",
        ),
        "capabilities": [
            ("Net position at a glance", "Due from Customers, Owed to Suppliers and the Net Position, headlined at the top."),
            ("Full receivables statement", "every customer balance with aging buckets and amount received in the period."),
            ("Full payables statement", "every supplier balance with aging, so upcoming obligations are clear."),
            ("Payment log", "a dated record of every receipt and payment in and out for the selected window."),
            ("Any period, one click to PDF", "choose a date range and print the entire statement for owners, banks or auditors."),
        ],
        "gains": [
            "An instant, honest answer to “can we meet our commitments this week?”",
            "Cash tied up in receivables set plainly against what you owe.",
            "A print-ready position statement whenever it is needed.",
        ],
        "images": [
            ("24-financial-position.png", "Financial Position: receivables vs. payables with the net position headlined."),
            ("22-report-outstanding.png", "Outstanding & Aging — receivables split into age buckets."),
            ("23-report-payables.png", "Supplier Payables — what you owe each supplier."),
        ],
    },
    {
        "cat": "TRANSACTIONS",
        "title": "Returns — Clean Reversals, No Manual Mess",
        "overview":
            "Returns are handled properly, not patched by hand. A sales return restores stock to the "
            "exact batches it was sold from and refunds the customer proportionally; a purchase "
            "return reverses stock and the supplier’s balance. Returns post immediately with the "
            "correct credit and debit notes, so inventory and ledgers stay perfectly in step without "
            "any manual adjustment.",
        "why": (
            "Why this matters",
            "In pharma, returns are constant — near-expiry stock comes back from pharmacies, and "
            "damaged or wrong goods go back to suppliers. Handled by hand, each one risks double-"
            "counting stock, refunding the wrong amount, or leaving a batch’s records inconsistent. "
            "WHMIS makes every reversal exact — back to the right batch, at the right value — so your "
            "inventory and your accounts never drift apart.",
        ),
        "capabilities": [
            ("Batch-accurate restock", "sales returns go back to the specific batches originally sold."),
            ("Proportional refunds", "the refund matches the value of the goods returned."),
            ("Supplier returns", "purchase returns reverse stock and the supplier balance cleanly."),
            ("Immediate posting", "returns settle at once with proper credit and debit notes."),
        ],
        "gains": [
            "Inventory and accounts that never drift after a return.",
            "No error-prone manual adjustments.",
            "A clear, documented trail for every return.",
        ],
        "images": [("17-returns.png", "Sales returns, tracked against the original invoice.")],
    },
    {
        "cat": "INSIGHTS",
        "title": "Reports & Analytics — Seventeen Reports, One Click to Export",
        "overview":
            "WHMIS ships with seventeen ready-made reports spanning Sales, Purchases, Inventory and "
            "Finance — the exact views a distributor needs to run and defend the business. Each one "
            "renders on screen with totals and, where useful, a chart, and each exports to Excel or "
            "PDF with a single click for sharing with owners, auditors or partners. There is nothing "
            "to build or configure; the insight is already there.",
        "why": (
            "Why this matters",
            "Data you cannot see is data you cannot act on. Pre-built, exportable reports mean you "
            "spend your time deciding, not assembling — and everyone, from the owner to the auditor, "
            "works from the same trusted figures.",
        ),
        "capabilities": [
            ("Sales insight", "registers, product and customer profitability, booker performance, and incentives given."),
            ("Purchase insight", "registers, supplier analysis, and bonus received versus bonus given."),
            ("Inventory insight", "stock position, stock movement, expiry, and slow- versus fast-movers."),
            ("Financial insight", "outstanding and aging, supplier payables, and a twelve-month profit trend."),
            ("Export anywhere", "one click to Excel or PDF on every report."),
        ],
        "gains": [
            "Immediate answers across every part of the business.",
            "Board-ready and audit-ready figures on demand.",
            "No spreadsheets to build or maintain.",
        ],
        "images": [
            ("18-reports-catalog.png", "The report catalogue, grouped by area."),
            ("19-report-profit.png", "A report with chart, totals and Excel / PDF export."),
            ("20-report-sales-register.png", "The sales register — one of seventeen built-in reports."),
        ],
    },
    {
        "cat": "PLATFORM",
        "title": "Workspace Tabs — Many Screens, One Window",
        "overview":
            "Distribution work is rarely one thing at a time — a call comes in while you are mid-"
            "invoice, and a report is needed while an order is half-entered. WHMIS lets you open "
            "sales, bookings, inventory and reports as tabs inside a single window and switch between "
            "them instantly. Each tab keeps its own data, filters and half-typed entries, so you "
            "never lose your place or your work when you jump between tasks.",
        "why": (
            "Why this matters",
            "Constantly reloading pages and re-entering filters wastes minutes that add up across a "
            "busy day. A tabbed workspace keeps your team in flow and their work safe.",
        ),
        "capabilities": [
            ("Multiple screens at once", "open several parts of the system side by side as tabs."),
            ("Instant switching", "move between tasks with no reload and no waiting."),
            ("State retained", "each tab remembers its data, filters and unsaved entries."),
            ("Familiar experience", "a browser-like feel your team already understands."),
        ],
        "gains": [
            "Higher productivity through uninterrupted multitasking.",
            "No lost work when switching between tasks.",
            "A modern, intuitive experience with almost no training.",
        ],
        "images": [("21-workspace.png", "Several screens open as tabs, each retaining its own state.")],
    },
    {
        "cat": "PLATFORM",
        "title": "Security & Access — The Right Access for Every Role",
        "overview":
            "WHMIS gives every team member exactly the access their job requires — and nothing more. "
            "Role-based permissions govern every sensitive action, from posting a sale to approving a "
            "booking, and roles such as booker are deliberately restricted so field staff can create "
            "orders but never touch billing or master pricing. Meanwhile, in-app notifications keep "
            "the right people informed about low stock, upcoming expiry, overdue balances and pending "
            "approvals.",
        "why": (
            "Why this matters",
            "Uncontrolled access is how prices get changed, discounts get abused and mistakes go "
            "unnoticed. Scoping access to roles protects both your data and your margin, and gives "
            "you a clear line of accountability.",
        ),
        "capabilities": [
            ("Role-based permissions", "every sensitive action is gated by role."),
            ("Least-privilege by design", "bookers and other roles see only what they need."),
            ("Proactive alerts", "notifications flag low stock, expiry, overdue balances and approvals awaiting action."),
            ("Accountability", "actions are tied to the user who performed them."),
        ],
        "gains": [
            "Sensitive functions protected from accidental or deliberate misuse.",
            "Clear accountability across the team.",
            "Important issues surfaced automatically, before they cost you.",
        ],
        "images": [("01-login.png", "Secure sign-in; access is scoped to each user’s role.")],
    },
    {
        "cat": "PLATFORM",
        "title": "Deployment & Localisation — Built to Fit How You Work",
        "overview":
            "WHMIS is delivered as a single, self-contained application that runs on standard cPanel "
            "web hosting — no specialist servers, no complex infrastructure and nothing for your "
            "staff to install. It is used from any modern browser on desktop or laptop, and it speaks "
            "your language commercially: prices in PKR and Pakistan tax fields (GST, NTN, STRN, CNIC) "
            "throughout, so your documents are compliant from the very first invoice. Your data lives "
            "in your own database, under your control.",
        "why": (
            "Why this matters",
            "Software that demands heavy infrastructure or foreign tax assumptions becomes a burden. "
            "WHMIS is designed to fit the way distributors in Pakistan already host, work and comply "
            "— so adoption is low-risk and low-cost.",
        ),
        "capabilities": [
            ("Simple hosting", "deploys to ordinary cPanel hosting; no special servers required."),
            ("Browser-based", "accessible from any modern browser, with nothing to install."),
            ("Pakistan-ready", "PKR currency and GST / NTN / STRN / CNIC fields built in."),
            ("Your data, your control", "everything is stored in your own database."),
        ],
        "gains": [
            "Low-cost, low-risk deployment on hosting you already understand.",
            "Compliant documents out of the box.",
            "Full ownership of your business data.",
        ],
        "images": [],
    },
]


# --------------------------------------------------------------------------- #
# Low-level docx helpers
# --------------------------------------------------------------------------- #
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def add_field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    instr = _el("w:instrText", **{"xml:space": "preserve"})
    instr.text = instruction
    run._r.append(instr)
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "separate"}))
    if placeholder:
        t = OxmlElement("w:t")
        t.text = placeholder
        run._r.append(t)
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))


def add_toc(document):
    p = document.add_paragraph()
    add_field(p, 'TOC \\o "1-1" \\h \\z \\u',
              "Right-click here and choose “Update Field” to build the contents.")


def spacer(document, pts=6):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(pts)
    return p


def para(document, text, size=11, color=None, after=8, before=0, align=None, italic=False, bold=False):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p


def eyebrow(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text.upper())
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = EMERALD


def heading_rule(document, title):
    h = document.add_heading(title, level=1)
    h.paragraph_format.space_before = Pt(2)
    h.paragraph_format.space_after = Pt(6)
    pbdr = _el("w:pBdr")
    pbdr.append(_el("w:bottom", **{"w:val": "single", "w:sz": "8", "w:space": "6", "w:color": "059669"}))
    h._p.get_or_add_pPr().append(pbdr)
    return h


def subhead(document, text):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = SLATE


def bullets(document, items, lead_color=SLATE):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        if isinstance(item, tuple):
            lead, desc = item
            r = p.add_run(lead + " — ")
            r.bold = True
            r.font.color.rgb = lead_color
            p.add_run(desc)
        else:
            p.add_run(item)


def callout(document, title, body):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(_el("w:shd", **{"w:val": "clear", "w:fill": CALLOUT_FILL}))
    # left accent + hairline box
    borders = _el("w:tcBorders")
    borders.append(_el("w:left", **{"w:val": "single", "w:sz": "24", "w:space": "0", "w:color": "059669"}))
    for edge in ("top", "bottom", "right"):
        borders.append(_el(f"w:{edge}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": "A7F3D0"}))
    tcPr.append(borders)
    mar = _el("w:tcMar")
    for m in ("top", "bottom", "left", "right"):
        mar.append(_el(f"w:{m}", **{"w:w": "120", "w:type": "dxa"}))
    tcPr.append(mar)

    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    r = p0.add_run(title)
    r.bold = True
    r.font.color.rgb = EMERALD_DK
    r.font.size = Pt(11)
    p1 = cell.add_paragraph()
    p1.paragraph_format.space_after = Pt(0)
    r = p1.add_run(body)
    r.font.size = Pt(10.5)
    r.font.color.rgb = SLATE
    spacer(document, 6)


def framed_image(document, path, width_in):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    borders = _el("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        borders.append(_el(f"w:{edge}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": RULE}))
    table._tbl.tblPr.append(borders)
    cell = table.rows[0].cells[0]
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].add_run().add_picture(path, width=Inches(width_in))


def caption(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY


# --------------------------------------------------------------------------- #
# Section renderer
# --------------------------------------------------------------------------- #
def render_section(document, num, sec, missing):
    eyebrow(document, f"Module {num:02d}  ·  {sec['cat']}")
    heading_rule(document, sec["title"])
    para(document, sec["overview"], after=8)
    callout(document, sec["why"][0], sec["why"][1])
    subhead(document, "Key capabilities")
    bullets(document, sec["capabilities"])
    subhead(document, "Business benefits")
    bullets(document, sec["gains"])
    imgs = sec.get("images", [])
    if imgs:
        spacer(document, 4)
        for name, cap in imgs:
            path = os.path.join(IMG, name)
            if os.path.exists(path):
                framed_image(document, path, CONTENT_WIDTH if len(imgs) == 1 else 6.0)
                caption(document, cap)
            else:
                missing.append(name)


# --------------------------------------------------------------------------- #
# Build
# --------------------------------------------------------------------------- #
def build():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for lvl, size, color in (("Title", 30, SLATE), ("Heading 1", 16, SLATE), ("Heading 2", 12.5, SLATE)):
        st = doc.styles[lvl]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True

    # ---------- Cover ----------
    for _ in range(5):
        doc.add_paragraph()
    para(doc, "WHMIS", size=48, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, bold=True)
    para(doc, "Pharmaceutical Distribution ERP", size=19, color=EMERALD,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    # emerald rule
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pbdr = _el("w:pBdr")
    pbdr.append(_el("w:bottom", **{"w:val": "single", "w:sz": "12", "w:space": "1", "w:color": "059669"}))
    line._p.get_or_add_pPr().append(pbdr)
    para(doc, "System Presentation", size=15, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

    for _ in range(7):
        doc.add_paragraph()
    para(doc, "Prepared for", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, CLIENT, size=18, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER, after=14, bold=True)
    para(doc, "A product of", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, VENDOR, size=13, color=EMERALD_DK, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, bold=True)

    for _ in range(4):
        doc.add_paragraph()
    para(doc, "Purchasing  •  Inventory  •  Sales  •  Bookings  •  Receivables  •  Reporting",
         size=10.5, color=LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    para(doc, date.today().strftime("%B %Y") + "   ·   Confidential",
         size=10, color=LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()

    # ---------- Confidentiality ----------
    doc.add_heading("Confidentiality Notice", level=1)
    para(doc,
         f"This document and the WHMIS platform it describes are the confidential and proprietary "
         f"property of {VENDOR}. It has been prepared exclusively for {CLIENT} for the sole purpose "
         f"of evaluating the system.",
         after=8)
    para(doc,
         "Its contents may not be copied, reproduced, distributed or disclosed to any third party, "
         "in whole or in part, without the prior written consent of " + VENDOR + ". By receiving "
         "this document, the recipient agrees to keep its contents confidential and to use them only "
         "for the intended evaluation.",
         after=8)
    para(doc, f"© {YEAR} {VENDOR}. All rights reserved.", size=10, color=GREY, after=0)
    doc.add_page_break()

    # ---------- Contents ----------
    doc.add_heading("Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # ---------- Executive summary ----------
    doc.add_heading("Executive Summary", level=1)
    for t in EXEC["paras"]:
        para(doc, t, after=8)
    callout(doc, EXEC["why"][0], EXEC["why"][1])
    subhead(doc, "What WHMIS delivers")
    bullets(doc, EXEC["gains"])
    doc.add_page_break()

    # ---------- Sections ----------
    missing = []
    for i, sec in enumerate(SECTIONS):
        render_section(doc, i + 1, sec, missing)
        if i != len(SECTIONS) - 1:
            doc.add_page_break()

    # ---------- Closing line ----------
    spacer(doc, 10)
    para(doc,
         f"Thank you. {VENDOR} would be glad to walk the team at {CLIENT} through any part of WHMIS "
         f"in a live, hands-on demonstration.",
         italic=True, color=SLATE, after=0)

    # ---------- Footer ----------
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run(f"Confidential  ·  WHMIS  ·  {VENDOR}  ·  Prepared for {CLIENT}  ·  Page ")
    add_field(fp, "PAGE")
    for run in fp.runs:
        run.font.size = Pt(7.5)
        run.font.color.rgb = GREY

    doc.save(OUT)
    print(f"Wrote {OUT}")
    if missing:
        print("WARNING missing images:", ", ".join(missing))


if __name__ == "__main__":
    build()
