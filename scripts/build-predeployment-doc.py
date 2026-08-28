#!/usr/bin/env python3
"""Build the WHMIS Pre-Deployment Requirements & Payment Schedule as a formatted .docx.

A premium, client-facing corporate commercial document between Virtual Wisdom
Technologies (Development / Technology Partner) and Master Pharmaceuticals (Client).

    python3 scripts/build-predeployment-doc.py
"""
import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTDIR = os.path.join(ROOT, "docs", "commercial")
OUT = os.path.join(OUTDIR, "WHMIS-PreDeployment-Requirements-and-Payment-Schedule.docx")

# Identity (exactly as specified for this document).
VENDOR = "Virtual Wisdom Technologies"
CLIENT = "Master Pharmaceuticals"
PROJECT = "WHMIS"
VERSION = "1.0"
DOC_DATE = "August 2026"
CLASSIFICATION = "Confidential – Commercial & Project Information"
DOC_TITLE = "WHMIS System — Pre-Deployment Requirements & Payment Schedule"

# Palette (slate + emerald corporate look, consistent with the WHMIS presentation).
SLATE = RGBColor(0x0F, 0x17, 0x2A)
EMERALD = RGBColor(0x05, 0x96, 0x69)
EMERALD_DK = RGBColor(0x04, 0x78, 0x57)
GREY = RGBColor(0x6B, 0x72, 0x80)
LIGHT = RGBColor(0x9C, 0xA3, 0xAF)
INK = RGBColor(0x1F, 0x29, 0x37)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
SLATE_HEX = "0F172A"
EMERALD_HEX = "059669"
CALLOUT_FILL = "ECFDF5"   # emerald-50
SNAP_FILL = "F0FDF4"      # emerald-50/green tint
ZEBRA_FILL = "F9FAFB"     # grey-50
TOTAL_FILL = "E5E7EB"     # grey-200
RULE = "D1D5DB"           # grey-300
CONTENT_WIDTH = 6.5       # inches, for 1-inch margins on Letter

# --------------------------------------------------------------------------- #
# Low-level helpers
# --------------------------------------------------------------------------- #
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def add_field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "begin"}))
    instr = _el("w:instrText", **{"xml:space": "preserve"})
    instr.text = instruction
    run._r.append(instr)
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "separate"}))
    if placeholder:
        t = OxmlElement("w:t")
        t.text = placeholder
        run._r.append(t)
    run._r.append(_el("w:fldChar", **{"w:fldCharType": "end"}))


def add_toc(document):
    p = document.add_paragraph()
    add_field(p, 'TOC \\o "1-1" \\h \\z \\u',
              "Right-click here and choose “Update Field” to build the contents.")


def spacer(document, pts=6):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(pts)
    return p


def para(document, text, size=11, color=None, after=8, before=0, align=None,
         italic=False, bold=False):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p


def runs_para(document, parts, after=8, before=0, align=None, size=11):
    """A paragraph built from (text, {bold,italic,color,size}) parts."""
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if align is not None:
        p.alignment = align
    for text, opts in parts:
        r = p.add_run(text)
        r.font.size = Pt(opts.get("size", size))
        r.bold = opts.get("bold", False)
        r.italic = opts.get("italic", False)
        if "color" in opts:
            r.font.color.rgb = opts["color"]
    return p


def heading_rule(document, title):
    h = document.add_heading(title, level=1)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    pbdr = _el("w:pBdr")
    pbdr.append(_el("w:bottom", **{"w:val": "single", "w:sz": "8", "w:space": "6", "w:color": EMERALD_HEX}))
    h._p.get_or_add_pPr().append(pbdr)
    return h


def subhead(document, text, before=10):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = SLATE
    return p


def bullets(document, items, lead_color=SLATE, after=4):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(after)
        if isinstance(item, tuple):
            lead, desc = item
            r = p.add_run(lead + " — ")
            r.bold = True
            r.font.color.rgb = lead_color
            p.add_run(desc)
        else:
            p.add_run(item)


def numbered(document, items, after=4):
    for item in items:
        p = document.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(after)
        p.add_run(item)


def _shade(cell, fill):
    cell._tc.get_or_add_tcPr().append(_el("w:shd", **{"w:val": "clear", "w:fill": fill}))


def _cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = _el("w:tcMar")
    for edge, w in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        mar.append(_el(f"w:{edge}", **{"w:w": str(w), "w:type": "dxa"}))
    tcPr.append(mar)


def _set_cell_text(cell, text, bold=False, color=None, size=10.5, align=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color


def table(document, headers, rows, widths, right_cols=(), total_row=None, after=8):
    """Clean commercial table: slate header, grey borders, zebra body, optional total row."""
    right_cols = set(right_cols)
    t = document.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    # borders
    borders = _el("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_el(f"w:{edge}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": RULE}))
    t._tbl.tblPr.append(borders)

    def _widths(row):
        for cell, w in zip(row.cells, widths):
            cell.width = Inches(w)

    # header
    hdr = t.rows[0]
    _widths(hdr)
    for i, (cell, text) in enumerate(zip(hdr.cells, headers)):
        _shade(cell, SLATE_HEX)
        _cell_margins(cell)
        _set_cell_text(cell, text, bold=True, color=WHITE, size=10.5,
                       align=WD_ALIGN_PARAGRAPH.RIGHT if i in right_cols else WD_ALIGN_PARAGRAPH.LEFT)
    # body
    for ri, row in enumerate(rows):
        r = t.add_row()
        _widths(r)
        for i, (cell, text) in enumerate(zip(r.cells, row)):
            if ri % 2 == 1:
                _shade(cell, ZEBRA_FILL)
            _cell_margins(cell)
            _set_cell_text(cell, str(text), size=10.5, color=INK,
                           align=WD_ALIGN_PARAGRAPH.RIGHT if i in right_cols else WD_ALIGN_PARAGRAPH.LEFT)
    # total
    if total_row is not None:
        r = t.add_row()
        _widths(r)
        for i, (cell, text) in enumerate(zip(r.cells, total_row)):
            _shade(cell, TOTAL_FILL)
            _cell_margins(cell)
            _set_cell_text(cell, str(text), bold=True, size=10.5, color=SLATE,
                           align=WD_ALIGN_PARAGRAPH.RIGHT if i in right_cols else WD_ALIGN_PARAGRAPH.LEFT)
    spacer(document, after)
    return t


def callout(document, title, body_parts, fill=CALLOUT_FILL, accent=EMERALD_HEX,
            title_color=EMERALD_DK):
    """Shaded left-accent box. body_parts: list of strings (paragraphs)."""
    t = document.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    cell = t.rows[0].cells[0]
    cell.width = Inches(CONTENT_WIDTH)
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(_el("w:shd", **{"w:val": "clear", "w:fill": fill}))
    borders = _el("w:tcBorders")
    borders.append(_el("w:left", **{"w:val": "single", "w:sz": "24", "w:space": "0", "w:color": accent}))
    for edge in ("top", "bottom", "right"):
        borders.append(_el(f"w:{edge}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": "A7F3D0"}))
    tcPr.append(borders)
    _cell_margins(cell, top=120, bottom=120, left=140, right=140)

    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(2)
    if title:
        r = p0.add_run(title)
        r.bold = True
        r.font.color.rgb = title_color
        r.font.size = Pt(11)
    for i, body in enumerate(body_parts):
        p = cell.add_paragraph() if (title or i > 0) else p0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(2 if (title or i > 0) else 0)
        r = p.add_run(body)
        r.font.size = Pt(10.5)
        r.font.color.rgb = INK
    spacer(document, 6)
    return t


# --------------------------------------------------------------------------- #
# Cover, parties, signatures
# --------------------------------------------------------------------------- #
def party_block(document, role, name, lines, name_color=SLATE):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(role.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = EMERALD
    _spc(r, 1.2)
    p2 = document.add_paragraph()
    p2.paragraph_format.space_after = Pt(1)
    r = p2.add_run(name)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = name_color
    for label, value in lines:
        pl = document.add_paragraph()
        pl.paragraph_format.space_after = Pt(0)
        rl = pl.add_run(label + "  ")
        rl.font.size = Pt(10)
        rl.font.color.rgb = GREY
        rv = pl.add_run(value)
        rv.font.size = Pt(10)
        rv.font.color.rgb = INK
    spacer(document, 8)


def _spc(run, pts):
    run._r.get_or_add_rPr().set("spc", str(int(pts * 100)))


def signature_block(document, org, name, designation):
    t = document.add_table(rows=1, cols=1)
    t.autofit = False
    cell = t.rows[0].cells[0]
    cell.width = Inches(CONTENT_WIDTH)
    _shade(cell, ZEBRA_FILL)
    borders = _el("w:tcBorders")
    for edge in ("top", "bottom", "left", "right"):
        borders.append(_el(f"w:{edge}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": RULE}))
    borders.insert(0, _el("w:left", **{"w:val": "single", "w:sz": "18", "w:space": "0", "w:color": EMERALD_HEX}))
    cell._tc.get_or_add_tcPr().append(borders)
    _cell_margins(cell, top=140, bottom=160, left=160, right=160)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("For " + org)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = SLATE

    for label, value in (("Name:", name), ("Designation:", designation),
                         ("Signature:", "__________________________"),
                         ("Date:", "__________________________")):
        pl = cell.add_paragraph()
        pl.paragraph_format.space_after = Pt(6)
        rl = pl.add_run(label + "  ")
        rl.font.size = Pt(10.5)
        rl.font.color.rgb = GREY
        rl.bold = True
        rv = pl.add_run(value)
        rv.font.size = Pt(10.5)
        rv.font.color.rgb = INK
    spacer(document, 8)


def snapshot_box(document, rows):
    """Highlighted commercial snapshot: rows = list of (label, value, emphasis?)."""
    t = document.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    cell = t.rows[0].cells[0]
    cell.width = Inches(CONTENT_WIDTH)
    _shade(cell, SNAP_FILL)
    borders = _el("w:tcBorders")
    for edge in ("top", "bottom", "right"):
        borders.append(_el(f"w:{edge}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": "A7F3D0"}))
    borders.insert(0, _el("w:left", **{"w:val": "single", "w:sz": "24", "w:space": "0", "w:color": EMERALD_HEX}))
    cell._tc.get_or_add_tcPr().append(borders)
    _cell_margins(cell, top=140, bottom=140, left=160, right=160)

    first = True
    for label, value, emph in rows:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(4 if not emph else 6)
        p.paragraph_format.space_before = Pt(0)
        if label:
            rl = p.add_run(label + "  ")
            rl.font.size = Pt(12 if emph else 10.5)
            rl.bold = True
            rl.font.color.rgb = SLATE
        if value:
            rv = p.add_run(value)
            rv.font.size = Pt(14 if emph else 10.5)
            rv.bold = emph
            rv.font.color.rgb = EMERALD_DK if emph else INK
    spacer(document, 8)


# --------------------------------------------------------------------------- #
# Build
# --------------------------------------------------------------------------- #
def build():
    os.makedirs(OUTDIR, exist_ok=True)
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    for lvl, size, color in (("Title", 30, SLATE), ("Heading 1", 15, SLATE), ("Heading 2", 12.5, SLATE)):
        st = doc.styles[lvl]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True

    # ---------------- Cover ----------------
    for _ in range(2):
        doc.add_paragraph()
    para(doc, CLASSIFICATION.upper(), size=9, color=EMERALD, after=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    para(doc, "WHMIS", size=52, color=SLATE, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, bold=True)
    para(doc, "Pharmaceutical Distribution ERP", size=15, color=EMERALD,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pbdr = _el("w:pBdr")
    pbdr.append(_el("w:bottom", **{"w:val": "single", "w:sz": "12", "w:space": "1", "w:color": EMERALD_HEX}))
    line._p.get_or_add_pPr().append(pbdr)
    para(doc, "Pre-Deployment Requirements & Payment Schedule", size=18, color=SLATE,
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2, bold=True)
    para(doc, f"Project / System: {PROJECT}     •     Version {VERSION}     •     {DOC_DATE}",
         size=10.5, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=18)

    # Parties (two-column table)
    pt = doc.add_table(rows=1, cols=2)
    pt.alignment = WD_TABLE_ALIGNMENT.CENTER
    pt.autofit = False
    left, right = pt.rows[0].cells
    left.width = Inches(3.25)
    right.width = Inches(3.25)
    for c in (left, right):
        _cell_margins(c, top=140, bottom=140, left=140, right=140)
        _shade(c, ZEBRA_FILL)
        b = _el("w:tcBorders")
        for edge in ("top", "bottom", "left", "right"):
            b.append(_el(f"w:{edge}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": RULE}))
        c._tc.get_or_add_tcPr().append(b)

    def _fill_party(cell, role, name, lines):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(role.upper()); r.bold = True; r.font.size = Pt(8.5); r.font.color.rgb = EMERALD; _spc(r, 1.0)
        p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(4)
        r = p2.add_run(name); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = SLATE
        for label, value in lines:
            pl = cell.add_paragraph(); pl.paragraph_format.space_after = Pt(1)
            rl = pl.add_run(label + " "); rl.font.size = Pt(9.5); rl.font.color.rgb = GREY
            rv = pl.add_run(value); rv.font.size = Pt(9.5); rv.font.color.rgb = INK

    _fill_party(left, "Development / Technology Partner", VENDOR, [
        ("Primary Contact:", "Muhammad Saad Mubeen"),
        ("Designation:", "Project Manager"),
        ("Email:", "pm@vwisdomtechnologies.com"),
        ("Phone:", "+92 318 5161571"),
    ])
    _fill_party(right, "Client", CLIENT, [
        ("Representative:", "Mr. Jamil Ahmed"),
        ("Designation:", "[Insert Designation]"),
        ("Email:", "[Insert Email]"),
        ("Phone:", "[Insert Contact Number]"),
        ("Address:", "[Insert Address]"),
    ])

    spacer(doc, 16)
    para(doc, f"Document Version {VERSION}   ·   {DOC_DATE}   ·   {CLASSIFICATION}",
         size=9.5, color=LIGHT, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.add_page_break()

    # ---------------- Confidentiality ----------------
    doc.add_heading("Confidentiality Notice", level=1)
    para(doc,
         f"This document is confidential and contains commercial and project information relating "
         f"to the {PROJECT} system. It is prepared by {VENDOR} for {CLIENT} for the sole purpose of "
         f"defining pre-deployment requirements, third-party costs and the associated payment "
         f"schedule.", after=8)
    para(doc,
         f"Its contents may not be copied, reproduced, distributed or disclosed to any third party, "
         f"in whole or in part, without the prior written consent of both parties. By receiving this "
         f"document, the recipient agrees to treat its contents as confidential.", after=8)
    para(doc, f"Classification: {CLASSIFICATION}.", size=10, color=GREY, italic=True, after=0)
    doc.add_page_break()

    # ---------------- Contents ----------------
    doc.add_heading("Contents", level=1)
    add_toc(doc)
    doc.add_page_break()

    # ---------------- 1. Purpose ----------------
    heading_rule(doc, "1. Purpose of the Document")
    para(doc,
         f"{VENDOR} has developed the existing Phase 1 of the {PROJECT} system for {CLIENT}. "
         f"{CLIENT} has decided to operate the solution as a cloud-based, internet-accessible system "
         f"so that authorized users can access it securely from anywhere over the internet.", after=8)
    para(doc, "The purpose of this document is to formally define:", after=4)
    bullets(doc, [
        "Requirements that must be completed before production deployment.",
        "Domain selection and registration.",
        "Hosting plan selection.",
        "Third-party domain and hosting charges.",
        "Phase 1 deployment payment.",
        "Remaining software development payment milestones.",
        "Production deployment responsibilities.",
        "Initial production go-live process.",
        "Post-deployment support and maintenance responsibility.",
        "Commercial dependencies affecting the implementation schedule.",
        "Next contractual and operational steps.",
    ])
    para(doc,
         "Detailed software development scope and future module delivery are covered separately "
         "under the WHMIS System – Phase-Wise Development & Deployment Plan.", before=4, after=8)
    para(doc,
         "Detailed contractual matters, support terms, maintenance terms, service levels, "
         "intellectual property, confidentiality, warranties, change requests and other "
         "legal/commercial provisions will subsequently be governed through the applicable Service "
         "Agreement / Software Development Agreement / Support & Maintenance terms.", after=8)

    # ---------------- 2. Production Deployment Approach ----------------
    heading_rule(doc, "2. Production Deployment Approach")
    para(doc,
         f"The {PROJECT} system will be deployed online so that authorized {CLIENT} users can access "
         f"it securely through the internet. The production deployment will be managed by {VENDOR}.", after=8)
    para(doc, f"{VENDOR} will handle all necessary technical deployment activities, including but not "
              f"limited to:", after=4)
    bullets(doc, [
        "Production server setup.", "Application deployment.",
        "Database deployment and configuration.", "Domain connection and DNS configuration.",
        "SSL/HTTPS configuration.", "Application environment configuration.",
        "Production security configuration.", "Backup configuration.",
        "Application storage configuration.", "Required production-level technical settings.",
        "Initial server/application monitoring setup.", "Deployment testing.",
        "Production verification.", "Administrative access configuration.",
        "Technical troubleshooting associated with deployment.",
    ])
    callout(doc, f"{CLIENT} only needs to approve three things",
            ["1. The selected domain.   2. The recommended hosting plan.   3. The applicable third-party charges.",
             f"Infrastructure-level technical details — CPU, RAM, operating system, database engine, "
             f"backup retention architecture, server configuration and similar parameters — are the "
             f"responsibility of {VENDOR}. {CLIENT} is not required to select or understand them."])

    # ---------------- 3. Domain Recommendations ----------------
    heading_rule(doc, "3. Domain Recommendations")
    para(doc,
         f"{VENDOR} will provide approximately 3–4 suitable domain recommendations for {CLIENT}. "
         f"The options below are placeholders and will be completed once domain research is finalised.", after=8)
    table(doc,
          ["Proposed Domain", "Availability", "First-Year Price", "Renewal Price", "Recommendation"],
          [
              ["[Domain Option 1]", "[To Be Checked]", "PKR [Insert Current Price]", "PKR [Insert Current Price]", "[Recommendation]"],
              ["[Domain Option 2]", "[To Be Checked]", "PKR [Insert Current Price]", "PKR [Insert Current Price]", "[Recommendation]"],
              ["[Domain Option 3]", "[To Be Checked]", "PKR [Insert Current Price]", "PKR [Insert Current Price]", "[Recommendation]"],
              ["[Domain Option 4]", "[To Be Checked]", "PKR [Insert Current Price]", "PKR [Insert Current Price]", "[Recommendation]"],
          ],
          widths=[1.9, 1.2, 1.2, 1.1, 1.1], right_cols=(2, 3))
    para(doc, "Please note the following regarding domain registration:", after=4)
    bullets(doc, [
        "Domain pricing is a third-party registrar charge.",
        "Domain cost is separate from the PKR 300,000 software development and deployment fee.",
        "Final domain availability must be verified immediately before purchase.",
        "Domain prices may change at any time.",
        "Promotional first-year prices may differ from renewal prices.",
        "Renewal rates are determined by the registrar.",
        "International vendor pricing may vary due to USD/PKR exchange-rate changes.",
        "Applicable taxes or registrar fees may affect the final price.",
        "Domain purchase will proceed after Master Pharmaceuticals approves one of the recommended options.",
    ])
    para(doc,
         f"Where possible, the final domain will be registered under {CLIENT}’ ownership, or under an "
         f"agreed ownership/account arrangement that ensures long-term client control.", before=4, after=8)

    # ---------------- 4. Recommended Hosting Plan ----------------
    heading_rule(doc, "4. Recommended Hosting Plan")
    para(doc,
         f"{VENDOR} will identify and recommend a suitable production hosting plan for the {PROJECT} "
         f"system. The recommendation is kept simple and commercial; the underlying infrastructure "
         f"engineering is handled by {VENDOR}.", after=8)
    table(doc,
          ["Hosting Plan", "Billing Period", "First-Year Cost", "Renewal Cost", "Recommendation"],
          [
              ["[Insert Provider] — [Insert Plan]", "[Annual / Monthly]", "PKR [Insert]", "PKR [Insert]",
               "Recommended for WHMIS Production Environment"],
          ],
          widths=[2.0, 1.1, 1.1, 1.0, 1.3], right_cols=(2, 3))
    para(doc,
         f"The recommended plan has been selected by {VENDOR} based on the application’s expected "
         f"operational needs. {CLIENT} is not required to make technical infrastructure decisions; "
         f"approval is only required for the selected plan and its cost.", after=8)

    # ---------------- 5. What VWT Will Manage ----------------
    heading_rule(doc, "5. What Virtual Wisdom Technologies Will Manage")
    para(doc,
         f"After approval of the hosting and domain, {VENDOR} will take responsibility for the "
         f"technical deployment and production configuration.", after=8)
    subhead(doc, "Production Deployment")
    bullets(doc, [
        "Hosting setup.", "Domain connection.", "SSL/HTTPS setup.", "Application deployment.",
        "Database setup.", "Required production configuration.", "User/admin setup.",
        "Deployment verification.", "Production testing.", "Go-live support.",
    ])
    subhead(doc, "Production Backup & Technical Configuration")
    para(doc,
         f"{VENDOR} will configure and manage the required backup arrangements according to the "
         f"selected hosting environment and project requirements. {CLIENT} is not required to define "
         f"the backup architecture.", after=6)
    subhead(doc, "Technical Server Management")
    para(doc,
         f"{VENDOR} will manage the technical configuration necessary for {PROJECT} to operate on the "
         f"selected hosting environment. The client will not be required to manage production "
         f"server-level settings unless otherwise mutually agreed.", after=8)

    # ---------------- 6. Third-Party Vendor Charges ----------------
    heading_rule(doc, "6. Third-Party Vendor Charges")
    para(doc, "The following costs are separate from the WHMIS software development and deployment "
              "charges:", after=4)
    bullets(doc, [
        "Domain registration.", "Domain renewal.", "Production hosting.", "Hosting renewal.",
        "Any separately approved premium third-party service introduced later.",
    ])
    para(doc, "These charges are payable based on the actual price charged by the applicable "
              "third-party provider.", before=4, after=8)
    callout(doc, "Domain and hosting costs are not included in the PKR 300,000 development and deployment fee.",
            ["Third-party vendor prices are subject to change according to vendor policies, "
             "promotional pricing, renewal pricing, applicable taxes, currency exchange-rate "
             "fluctuations and other charges imposed by the respective service provider. "
             f"{VENDOR} does not control third-party pricing or renewal rates."])

    # ---------------- 7. Total Development & Deployment Charges ----------------
    heading_rule(doc, "7. Total Development & Deployment Charges")
    runs_para(doc, [("Total Agreed Software Development & Deployment Value:  ", {"bold": True, "size": 12, "color": SLATE}),
                    ("PKR 300,000", {"bold": True, "size": 13, "color": EMERALD_DK})], after=8)
    para(doc, "The amount is divided across four project phases:", after=6)
    table(doc,
          ["Phase", "Scope", "Amount"],
          [
              ["Phase 1", "Existing WHMIS Core System – Production Deployment", "PKR 100,000"],
              ["Phase 2", "HR, Employees & Payroll Module", "PKR 60,000"],
              ["Phase 3", "Medical Representatives, Doctors & Field Visit Management", "PKR 70,000"],
              ["Phase 4", "Doctor Activities, Financial Records, Samples, Integration & Final Enhancements", "PKR 70,000"],
          ],
          widths=[1.0, 4.2, 1.3], right_cols=(2,),
          total_row=["Total", "Complete Agreed Development & Deployment Scope", "PKR 300,000"])
    runs_para(doc, [("Note:  ", {"bold": True, "color": SLATE}),
                    ("Domain registration, hosting and any approved third-party services are charged separately.",
                     {"bold": True})], after=8)

    # ---------------- 8. Payment Schedule ----------------
    heading_rule(doc, "8. Payment Schedule")

    def milestone(title, amount, scope_label, scope, trigger_label, trigger, extra_intro=None, extra=None):
        subhead(doc, title, before=8)
        runs_para(doc, [("Amount:  ", {"bold": True, "color": SLATE}),
                        (amount, {"bold": True, "color": EMERALD_DK, "size": 12})], after=4)
        runs_para(doc, [(scope_label + ":  ", {"bold": True, "color": SLATE}), (scope, {})], after=4)
        runs_para(doc, [(trigger_label + ":  ", {"bold": True, "color": SLATE}), (trigger, {})], after=4)
        if extra_intro:
            para(doc, extra_intro, after=4)
        if extra:
            bullets(doc, extra)

    milestone("Milestone 1 – Phase 1", "PKR 100,000",
              "Purpose",
              "Deployment of the currently completed Phase 1 WHMIS system into the production cloud environment.",
              "Payment Due",
              "Before commencement/finalization of production deployment.",
              extra_intro="Phase 1 production deployment will proceed after:",
              extra=[
                  "Approval of this document.",
                  "Receipt of PKR 100,000 Phase 1 payment.",
                  "Approval of the selected domain.",
                  "Approval of the recommended hosting plan.",
                  "Payment/authorization of third-party domain and hosting charges.",
                  "Provision of any essential business information required for production setup.",
              ])
    milestone("Milestone 2 – Phase 2", "PKR 60,000",
              "Scope", "HR, Employee and Payroll Module.",
              "Payment Trigger",
              "Before commencement of Phase 2 development after completion/deployment of Phase 1 and authorization to proceed.")
    milestone("Milestone 3 – Phase 3", "PKR 70,000",
              "Scope", "Medical Representative, Doctor and Field Visit Management.",
              "Payment Trigger",
              "Before commencement of Phase 3 after review/acceptance of Phase 2 and authorization to proceed.")
    milestone("Milestone 4 – Phase 4", "PKR 70,000",
              "Scope", "Doctor Activity & Financial Records, Sample Management, Integration and Final Enhancements.",
              "Payment Trigger",
              "Before commencement of Phase 4 after review/acceptance of Phase 3 and authorization to proceed.")
    spacer(doc, 6)

    # ---------------- 9. Payment Summary ----------------
    heading_rule(doc, "9. Payment Summary")
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    borders = _el("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_el(f"w:{edge}", **{"w:val": "single", "w:sz": "4", "w:space": "0", "w:color": RULE}))
    t._tbl.tblPr.append(borders)
    summary = [
        ("Phase 1", "PKR 100,000", False),
        ("Phase 2", "PKR 60,000", False),
        ("Phase 3", "PKR 70,000", False),
        ("Phase 4", "PKR 70,000", False),
        ("Total Software Development & Deployment", "PKR 300,000", True),
        ("Domain", "Separate – Actual Vendor Cost", False),
        ("Hosting", "Separate – Actual Vendor Cost", False),
    ]
    for i, (label, value, emph) in enumerate(summary):
        row = t.add_row()
        row.cells[0].width = Inches(4.3)
        row.cells[1].width = Inches(2.2)
        for j, (cell, text) in enumerate(zip(row.cells, (label, value))):
            _cell_margins(cell)
            if emph:
                _shade(cell, TOTAL_FILL)
            elif i % 2 == 1:
                _shade(cell, ZEBRA_FILL)
            _set_cell_text(cell, text, bold=emph, size=10.5,
                           color=SLATE if emph else INK,
                           align=WD_ALIGN_PARAGRAPH.RIGHT if j == 1 else WD_ALIGN_PARAGRAPH.LEFT)
    spacer(doc, 8)

    # ---------------- 10. Phase 1 Pre-Deployment & Go-Live Schedule ----------------
    heading_rule(doc, "10. Phase 1 Pre-Deployment & Go-Live Schedule")
    para(doc,
         "Assuming commercial approval, Phase 1 payment and domain/hosting approval are completed by "
         "24 August 2026, the proposed schedule is:", after=6)
    table(doc,
          ["Date", "Activity"],
          [
              ["24 August 2026", "Commercial approval and Phase 1 payment confirmation"],
              ["24–25 August 2026", "Domain and hosting purchase/provisioning"],
              ["25–27 August 2026", f"Production environment setup by {VENDOR}"],
              ["26–28 August 2026", "WHMIS application and database deployment"],
              ["28 August 2026", "Internal production verification"],
              ["29 August 2026", "Proposed Phase 1 production go-live"],
          ],
          widths=[1.9, 4.6])
    callout(doc, "These are proposed target dates.",
            ["The dates above depend on timely payment, approvals and availability of required "
             "information. If the approval/payment date changes, the deployment timeline will move "
             "accordingly. 29 August 2026 is a proposed go-live target, not an unconditional "
             "contractual deadline."],
            fill=ZEBRA_FILL, accent="9CA3AF", title_color=SLATE)

    # ---------------- 11. Responsibilities of VWT ----------------
    heading_rule(doc, "11. Responsibilities of Virtual Wisdom Technologies")
    para(doc, f"{VENDOR} will be responsible for:", after=4)
    bullets(doc, [
        "Recommending suitable domain options.", "Providing current domain pricing.",
        "Recommending the production hosting plan.", "Providing current hosting pricing.",
        "Purchasing/configuring approved domain and hosting after client authorization/payment.",
        "Production server setup.", "Production application deployment.",
        "Production database deployment.", "Domain/DNS configuration.", "SSL/HTTPS setup.",
        "Backup configuration.", "Production security configuration.", "Application settings.",
        "User/admin configuration as required.", "Deployment QA.", "Production testing.",
        "Initial go-live technical support.", "Resolving deployment-related technical issues.",
        "Maintaining technical production configuration under the agreed support arrangement.",
    ])

    # ---------------- 12. Responsibilities of Master Pharmaceuticals ----------------
    heading_rule(doc, "12. Responsibilities of Master Pharmaceuticals")
    para(doc, f"{CLIENT} will be responsible for:", after=4)
    bullets(doc, [
        "Reviewing and approving one of the proposed domains.",
        "Approving the recommended hosting plan.",
        "Paying the applicable domain and hosting vendor charges.",
        "Paying the applicable software-development milestone.",
        "Providing required authorized representative details.",
        "Providing relevant company/business information required for WHMIS.",
        "Providing required users/operational data where applicable.",
        "Reviewing delivered functionality.",
        "Providing timely and consolidated feedback.",
        "Providing UAT/operational approval as required.",
    ])
    para(doc,
         f"{CLIENT} is not responsible for technical server architecture, backups, SSL setup, "
         f"database administration or production deployment configuration — these are handled by "
         f"{VENDOR}.", before=4, after=8)

    # ---------------- 13. Post-Deployment Support & Maintenance ----------------
    heading_rule(doc, "13. Post-Deployment Support & Maintenance")
    para(doc,
         f"{VENDOR} will also provide support and maintenance services for the deployed {PROJECT} "
         f"system according to the mutually agreed support arrangement. The support and maintenance "
         f"function may include:", after=4)
    bullets(doc, [
        "Production application support.", "Technical troubleshooting.",
        "Correction of software defects within the supported deployed scope.",
        "Deployment-related issue resolution.", "Backup monitoring/management where applicable.",
        "Production environment maintenance.", "Basic security/configuration maintenance.",
        "Application updates relating to the agreed system.",
        "Assistance with system operational issues.",
        "Restoration/support activities where technically required and available.",
        "Coordination with the hosting/vendor where necessary.",
        "Reasonable post-deployment technical assistance.",
    ])
    subhead(doc, "Included Support / Maintenance")
    para(doc,
         "Routine support and maintenance of the agreed deployed WHMIS system according to the final "
         "support arrangement.", after=6)
    subhead(doc, "Additional Development")
    para(doc,
         "New modules, substantial workflow changes, integrations, features or other enhancements "
         "outside the agreed development scope are not maintenance and will be treated as separate "
         "change requests or future development phases.", after=6)
    callout(doc, "Support & Maintenance Commercial Terms",
            ["To be documented in the final Service Agreement / Support & Maintenance Schedule."],
            fill=ZEBRA_FILL, accent="9CA3AF", title_color=SLATE)

    # ---------------- 14. Change Requests ----------------
    heading_rule(doc, "14. Change Requests")
    para(doc,
         "PKR 300,000 represents the currently agreed development/deployment scope defined in the "
         "approved Phase-Wise Development & Deployment Plan. Any substantial new requirement will be "
         "treated as a Change Request and evaluated for:", after=4)
    bullets(doc, [
        "Functional impact.", "Technical impact.", "Development effort.", "Timeline impact.",
        "Commercial impact.",
    ])
    para(doc,
         "No substantial additional development should be assumed to be included without written "
         "approval.", before=4, after=8)

    # ---------------- 15. Taxes and External Charges ----------------
    heading_rule(doc, "15. Taxes and External Charges")
    para(doc,
         "Applicable taxes, withholding deductions, banking charges or legally applicable government "
         "levies will be handled according to applicable law and the final invoice/payment "
         "arrangement.", after=8)

    # ---------------- 16. Commercial Snapshot ----------------
    heading_rule(doc, "16. Commercial Snapshot")
    snapshot_box(doc, [
        ("Project:", PROJECT, False),
        ("Client:", CLIENT, False),
        ("Technology Partner:", VENDOR, False),
        ("Total Development & Deployment Charges:", "PKR 300,000", True),
        ("Phase 1:", "PKR 100,000        Phase 2:  PKR 60,000        Phase 3:  PKR 70,000        Phase 4:  PKR 70,000", False),
        ("Domain:", "Separate third-party vendor cost", False),
        ("Hosting:", "Separate third-party vendor cost", False),
        ("Proposed Phase 1 Go-Live:", "29 August 2026, subject to completion of prerequisites.", False),
        ("Support & Maintenance:", f"To be provided by {VENDOR} under the agreed support arrangement.", False),
    ])

    # ---------------- 17. Next Steps ----------------
    heading_rule(doc, "17. Next Steps")
    numbered(doc, [
        "Approval of this Pre-Deployment Requirements & Payment Schedule.",
        "Review of recommended domain options.",
        "Selection and approval of domain.",
        "Approval of recommended hosting plan.",
        "Payment of domain and hosting vendor charges.",
        "Payment of Phase 1 – PKR 100,000.",
        f"Production deployment by {VENDOR}.",
        "Phase 1 go-live.",
        "Commencement of Phase 2 according to the separate Phase-Wise Development & Deployment Plan.",
        "Subsequent execution/finalization of Service Agreement and Support & Maintenance terms.",
    ])
    spacer(doc, 6)

    # ---------------- 18. Acknowledgement & Sign-Off ----------------
    heading_rule(doc, "18. Acknowledgement & Sign-Off")
    para(doc,
         "By acknowledging this document, both parties confirm their understanding of the "
         "pre-deployment requirements, third-party domain and hosting costs, development payment "
         "milestones and proposed deployment process. Detailed software scope is governed by the "
         "approved Phase-Wise Development & Deployment Plan, while broader contractual and support "
         "terms will be covered under the applicable Service Agreement and Support & Maintenance "
         "arrangements.", after=10)
    signature_block(doc, VENDOR, "Muhammad Azhar", "[Insert]")
    signature_block(doc, CLIENT, "[Insert]", "[Insert]")

    # ---------------- Header / Footer ----------------
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run(f"{DOC_TITLE}  ·  v{VERSION}")
    hr.font.size = Pt(8)
    hr.font.color.rgb = LIGHT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(f"{CLASSIFICATION}     ·     {VENDOR}     ·     {PROJECT}     ·     Page ")
    fr.font.size = Pt(8)
    fr.font.color.rgb = GREY
    add_field(fp, "PAGE")
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GREY

    # Suppress header/footer on the cover (first page).
    section.different_first_page_header_footer = True

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
